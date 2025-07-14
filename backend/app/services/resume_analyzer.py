import os
import re
import io
import json
from werkzeug.utils import secure_filename
from PyPDF2 import PdfReader
import docx
from datetime import datetime
from .llm_service import get_llm_response

# Constantes pour les types de fichiers supportés
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'doc', 'txt'}

# Template de prompt pour l'analyse du CV
RESUME_ANALYSIS_PROMPT = """
Tu es un expert en recrutement chargé d'analyser un CV pour le poste de {job_role}.

Voici le contenu textuel du CV à analyser:
```
{resume_text}
```

Effectue une analyse approfondie de ce CV et fournis un rapport détaillé sur les éléments suivants:

1. Résumé des qualifications clés
2. Compétences techniques (technologies, langages, outils)
3. Compétences non techniques (soft skills)
4. Expérience professionnelle pertinente pour le poste
5. Formation et certifications
6. Adéquation globale avec le poste (score sur 10 et justification)
7. Forces principales du candidat
8. Éventuelles lacunes ou domaines à approfondir
9. Questions recommandées pour l'entretien, basées sur ce CV

Fournis ta réponse au format JSON avec la structure suivante:
{{
  "resume_summary": "Résumé concis du profil",
  "technical_skills": ["compétence1", "compétence2", ...],
  "soft_skills": ["compétence1", "compétence2", ...],
  "relevant_experience": [
    {{
      "position": "Titre du poste",
      "company": "Entreprise",
      "duration": "Durée",
      "highlights": ["point clé 1", "point clé 2", ...]
    }},
    ...
  ],
  "education": [
    {{
      "degree": "Diplôme/Formation",
      "institution": "Institution",
      "year": "Année",
      "relevance": "Pertinence pour le poste (haute/moyenne/faible)"
    }},
    ...
  ],
  "fit_score": 7.5,
  "fit_justification": "Explication du score d'adéquation",
  "strengths": ["force 1", "force 2", ...],
  "gaps": ["lacune 1", "lacune 2", ...],
  "recommended_questions": [
    {{
      "question": "Question d'entretien",
      "rationale": "Pourquoi poser cette question"
    }},
    ...
  ],
  "keywords": ["mot-clé1", "mot-clé2", ...] // mots-clés pertinents extraits du CV
}}
"""

CV_IMPROVEMENT_PROMPT = """
Tu es un expert en recrutement et en optimisation de CV. Analyse ce CV et fais des recommandations 
d'amélioration pour {improvement_target}.

Voici le contenu textuel du CV à analyser:
{resume_text}
{additional_context}

Fournis une analyse approfondie et des suggestions d'amélioration concrètes pour ce CV.
Ta réponse doit suivre le format JSON suivant:
{{
  "analysis": {{
    "strengths": ["point fort 1", "point fort 2", ...],
    "weaknesses": ["point faible 1", "point faible 2", ...],
    "missing_elements": ["élément manquant 1", "élément manquant 2", ...],
    "format_issues": ["problème de format 1", "problème de format 2", ...]
  }},
  "recommendations": {{
    "content_improvements": [
      {{
        "section": "section du CV (ex: Expérience professionnelle, Formation, etc.)",
        "current_content": "contenu actuel (résumé)",
        "suggested_improvement": "suggestion d'amélioration détaillée",
        "rationale": "justification de cette suggestion"
      }},
      ...
    ],
    "structure_improvements": [
      {{
        "recommendation": "recommandation spécifique",
        "importance": "élevée/moyenne/faible",
        "details": "détails de la recommandation"
      }},
      ...
    ],
    "keyword_optimization": ["mot-clé 1", "mot-clé 2", ...],
    "ats_compatibility_tips": ["conseil 1", "conseil 2", ...] 
  }},
  "improved_summary": "Un résumé professionnel amélioré pour le début du CV"
}}
"""

CV_TEMPLATES = {
    "american": {
        "sections": [
            "PROFESSIONAL SUMMARY",
            "SKILLS",
            "EXPERIENCE",
            "EDUCATION",
            "CERTIFICATIONS"
        ],
        "format_guidelines": {
            "length": "1-2 pages maximum",
            "font": "Arial or Times New Roman, 10-12pt",
            "bullet_points": "Start with action verbs, focus on achievements",
            "sections_order": "Summary → Skills → Experience → Education",
            "contact_info": "Name, phone, email, LinkedIn (optional: location)",
            "design": "Clean and minimal, limited use of color"
        }
    },
    "canadian": {
        "sections": [
            "PROFESSIONAL SUMMARY",
            "SKILLS",
            "EXPERIENCE",
            "EDUCATION",
            "VOLUNTEER EXPERIENCE",
            "LANGUAGES"
        ],
        "format_guidelines": {
            "length": "1-2 pages maximum",
            "font": "Arial or Calibri, 11-12pt",
            "bullet_points": "Focus on achievements and transferable skills",
            "sections_order": "Summary → Skills → Experience → Education → Languages",
            "contact_info": "Name, phone, email, LinkedIn, location (city, province)",
            "design": "Clean, professional, minimal formatting"
        }
    },
    "linkedin": {
        "sections": [
            "ABOUT",
            "EXPERIENCE",
            "EDUCATION",
            "SKILLS",
            "RECOMMENDATIONS",
            "ACCOMPLISHMENTS"
        ],
        "format_guidelines": {
            "length": "Concise sections, scannable format",
            "headline": "Job title + value proposition (limited to 120 characters)",
            "about": "First-person narrative, 3-5 paragraphs maximum",
            "experience": "Focus on achievements with metrics, use 3-5 bullet points per role",
            "skills": "List technical and soft skills (maximum 50)",
            "profile_photo": "Professional headshot recommended",
            "keywords": "Industry-specific keywords throughout profile"
        }
    }
}

def allowed_file(filename):
    """
    Vérifie si le fichier a une extension autorisée
    """
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(pdf_file):
    """
    Extrait le texte d'un fichier PDF
    """
    pdf_reader = PdfReader(pdf_file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text() + "\n"
    return text

def extract_text_from_docx(docx_file):
    """
    Extrait le texte d'un fichier DOCX
    """
    doc = docx.Document(docx_file)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

def extract_text_from_file(file):
    """
    Extrait le texte d'un fichier selon son type
    """
    filename = secure_filename(file.filename)
    file_extension = filename.rsplit('.', 1)[1].lower() if '.' in filename else ""
    
    # Réinitialiser la position du curseur au début du fichier
    file.seek(0)
    
    if file_extension == 'pdf':
        return extract_text_from_pdf(file)
    elif file_extension == 'docx':
        return extract_text_from_docx(file)
    elif file_extension == 'txt':
        return file.read().decode('utf-8')
    else:
        raise ValueError(f"Format de fichier non pris en charge: {file_extension}")

def preprocess_resume_text(text):
    """
    Prétraite le texte du CV pour améliorer l'analyse
    """
    # Suppression des espaces multiples
    text = re.sub(r'\s+', ' ', text)
    
    # Suppression des caractères spéciaux superflus
    text = re.sub(r'[^\w\s.,:;@()[\]{}\'\"&\-/]', '', text)
    
    # Normalisation des sauts de ligne
    text = re.sub(r'\n+', '\n', text)
    
    return text.strip()

def extract_contact_info(text):
    """
    Extrait les informations de contact du texte du CV
    """
    contact_info = {
        'email': None,
        'phone': None,
        'linkedin': None
    }
    
    # Extraction de l'email
    email_pattern = r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+'
    email_matches = re.findall(email_pattern, text)
    if email_matches:
        contact_info['email'] = email_matches[0]
    
    # Extraction du numéro de téléphone (formats français et internationaux)
    phone_patterns = [
        r'(?:\+33|0)(?:\s*\d){9}',  # Format français
        r'(?:\+\d{1,3}\s*)?\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4}'  # Format international
    ]
    
    for pattern in phone_patterns:
        phone_matches = re.findall(pattern, text)
        if phone_matches:
            contact_info['phone'] = phone_matches[0]
            break
    
    # Extraction du profil LinkedIn
    linkedin_patterns = [
        r'linkedin\.com/in/[\w-]+',
        r'linkedin\.com/profile/[\w-]+'
    ]
    
    for pattern in linkedin_patterns:
        linkedin_matches = re.findall(pattern, text)
        if linkedin_matches:
            contact_info['linkedin'] = linkedin_matches[0]
            break
    
    return contact_info

def extract_experience_years(resume_text):
    """
    Extrait les années d'expérience professionnelle d'un CV avec une meilleure précision
    
    Args:
        resume_text (str): Le texte du CV
        
    Returns:
        float: Années d'expérience (avec décimales pour les mois)
    """
    import re
    from datetime import datetime
    import calendar
    
    experience_years = []
    total_experience = 0
    
    # 1. Recherche d'une mention explicite des années d'expérience
    experience_patterns = [
        # Français
        r'(\d+)[+\-]?\s*(?:an(?:s|née(?:s)?)?\s+d\'?(?:expérience))',
        r'expérience\s+(?:de|:)?\s*(\d+)[+\-]?\s*an(?:s|née(?:s)?)?',
        r'(?:minimum|au moins)\s+(\d+)\s*an(?:s|née(?:s)?)?',
        r'depuis\s+(\d{4})',
        # Anglais
        r'(\d+)[+\-]?\s*(?:year(?:s)?\s+(?:of\s+)?experience)',
        r'experience\s+(?:of|:)?\s*(\d+)[+\-]?\s*year(?:s)?',
        r'(?:minimum|at least)\s+(\d+)\s*year(?:s)?',
        r'since\s+(\d{4})'
    ]
    
    for pattern in experience_patterns:
        matches = re.finditer(pattern, resume_text.lower())
        for match in matches:
            try:
                if pattern == r'depuis\s+(\d{4})' or pattern == r'since\s+(\d{4})':
                    # Calculer les années depuis l'année mentionnée
                    year = int(match.group(1))
                    current_year = datetime.now().year
                    experience_years.append(current_year - year)
                else:
                    # Extraction directe du nombre d'années
                    years = int(match.group(1))
                    experience_years.append(years)
            except (IndexError, ValueError) as e:
                print(f"Erreur lors de l'extraction des années d'expérience: {str(e)}")
                continue
    
    # 2. Identifier les emplois et leurs durées 
    job_entries = []
    
    # A. Détection des dates avec format mois-année à mois-année
    month_year_pattern = r'((?:janvier|février|mars|avril|mai|juin|juillet|août|septembre|octobre|novembre|décembre|jan|fév|mar|avr|mai|juin|juil|août|sept|oct|nov|déc|january|february|march|april|may|june|july|august|september|october|november|december|jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\.?\s+\d{4})\s*(?:[-–]\s*|\s+(?:à|au|to|until|jusqu\'(?:à|au))\s+)((?:janvier|février|mars|avril|mai|juin|juillet|août|septembre|octobre|novembre|décembre|jan|fév|mar|avr|mai|juin|juil|août|sept|oct|nov|déc|january|february|march|april|may|june|july|august|september|october|november|december|jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\.?\s+\d{4}|(?:aujourd\'hui|present|actuel|maintenant|current|now|today|présent))'
    
    month_year_matches = re.finditer(month_year_pattern, resume_text.lower())
    for match in month_year_matches:
        job_entries.append((match.group(1), match.group(2)))
    
    # B. Détection des dates avec jours précis (pour les stages courts)
    day_month_year_pattern = r'(\d{1,2}\s+(?:janvier|février|mars|avril|mai|juin|juillet|août|septembre|octobre|novembre|décembre|jan|fév|mar|avr|mai|juin|juil|août|sept|oct|nov|déc|january|february|march|april|may|june|july|august|september|october|november|december|jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\.?\s+\d{4})\s*(?:[-–]\s*|\s+(?:à|au|to|until|jusqu\'(?:à|au))\s+)(\d{1,2}\s+(?:janvier|février|mars|avril|mai|juin|juillet|août|septembre|octobre|novembre|décembre|jan|fév|mar|avr|mai|juin|juil|août|sept|oct|nov|déc|january|february|march|april|may|june|july|august|september|october|november|december|jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\.?\s+\d{4})'
    
    day_month_year_matches = re.finditer(day_month_year_pattern, resume_text.lower())
    for match in day_month_year_matches:
        job_entries.append((match.group(1), match.group(2)))
    
    # C. Format "Mois-Mois YYYY" (ex: "Août-septembre 2023")
    month_to_month_pattern = r'((?:janvier|février|mars|avril|mai|juin|juillet|août|septembre|octobre|novembre|décembre|jan|fév|mar|avr|mai|juin|juil|août|sept|oct|nov|déc|january|february|march|april|may|june|july|august|september|october|november|december|jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\.?)[- ](?:à|au|to|until|jusqu\'(?:à|au))?\s*((?:janvier|février|mars|avril|mai|juin|juillet|août|septembre|octobre|novembre|décembre|jan|fév|mar|avr|mai|juin|juil|août|sept|oct|nov|déc|january|february|march|april|may|june|july|august|september|october|november|december|jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\.?)\s+(\d{4})'
    
    month_to_month_matches = re.finditer(month_to_month_pattern, resume_text.lower())
    for match in month_to_month_matches:
        # Convertir les mois en nombres pour calculer la durée
        months_fr = {'janvier': 1, 'février': 2, 'mars': 3, 'avril': 4, 'mai': 5, 'juin': 6, 
                    'juillet': 7, 'août': 8, 'septembre': 9, 'octobre': 10, 'novembre': 11, 'décembre': 12,
                    'jan': 1, 'fév': 2, 'mar': 3, 'avr': 4, 'mai': 5, 'juin': 6, 
                    'juil': 7, 'août': 8, 'sept': 9, 'oct': 10, 'nov': 11, 'déc': 12}
        
        months_en = {'january': 1, 'february': 2, 'march': 3, 'april': 4, 'may': 5, 'june': 6, 
                    'july': 7, 'august': 8, 'september': 9, 'october': 10, 'november': 11, 'december': 12,
                    'jan': 1, 'feb': 2, 'mar': 3, 'apr': 4, 'may': 5, 'jun': 6, 
                    'jul': 7, 'aug': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dec': 12}
        
        all_months = {**months_fr, **months_en}
        
        start_month = match.group(1).lower()
        end_month = match.group(2).lower()
        year = int(match.group(3))
        
        # Obtenir le numéro du mois
        start_month_num = None
        for m in all_months:
            if m in start_month:
                start_month_num = all_months[m]
                break
        
        end_month_num = None
        for m in all_months:
            if m in end_month:
                end_month_num = all_months[m]
                break
        
        if start_month_num and end_month_num:
            start_date = f"{start_month_num}/01/{year}"
            end_date = f"{end_month_num}/28/{year}"
            job_entries.append((start_date, end_date))
    
    # D. Années (YYYY-YYYY) - mais uniquement pour les expériences professionnelles
    year_to_year_pattern = r'(\d{4})\s*[àa\-]\s*(20\d{2}|aujourd\'hui|present|actuel|maintenant|current|now|today|présent)'
    
    year_matches = re.finditer(year_to_year_pattern, resume_text.lower())
    for match in year_matches:
        # Ignorer les correspondances qui semblent être des périodes d'éducation
        context_before = resume_text[max(0, match.start() - 50):match.start()].lower()
        if any(edu_term in context_before for edu_term in ['formation', 'education', 'études', 'diplôme', 'baccalauréat', 'licence', 'master', 'dut']):
            continue
            
        job_entries.append((match.group(1), match.group(2)))
    
    # 3. Calculer la durée en mois pour chaque entrée, puis convertir en années
    from datetime import datetime
    
    for start, end in job_entries:
        try:
            # Vérifier si l'entrée contient des mots indiquant un stage
            is_internship = False
            job_context = ""
            
            # Extraire le contexte autour de cette entrée pour déterminer s'il s'agit d'un stage
            try:
                start_pos = resume_text.lower().find(str(start).lower())
                end_pos = resume_text.lower().find(str(end).lower())
                if start_pos >= 0 and end_pos >= 0:
                    job_context = resume_text[max(0, start_pos - 100):min(len(resume_text), end_pos + 100)]
            except:
                # Si on ne peut pas trouver le contexte, on continue avec une valeur par défaut
                pass
            
            if re.search(r'stage|stagiaire|intern|internship', job_context.lower()):
                is_internship = True
                
            # A. Format année-année
            if re.match(r'\d{4}$', str(start)) and (re.match(r'\d{4}$', str(end)) or end.lower() in ['aujourd\'hui', 'present', 'actuel', 'maintenant', 'current', 'now', 'today', 'présent']):
                start_year = int(start)
                if re.match(r'\d{4}$', str(end)):
                    end_year = int(end)
                else:
                    end_year = datetime.now().year
                
                # Vérifier la cohérence des années
                if 1950 <= start_year <= datetime.now().year and start_year <= end_year:
                    # Convertir en mois (approximation)
                    months = (end_year - start_year) * 12
                    
                    # Si c'est un stage, limiter la durée maximum à 6 mois sauf indication contraire
                    if is_internship and months > 6:
                        months = 6
                        
                    experience_years.append(months / 12)
            
            # B. Formats de date plus précis
            else:
                try:
                    # Essayer différents formats de date
                    date_formats = [
                        '%m/%d/%Y',  # MM/DD/YYYY
                        '%d/%m/%Y',  # DD/MM/YYYY
                        '%B %Y',     # Month YYYY
                        '%b %Y',     # Mon YYYY
                        '%d %B %Y',  # DD Month YYYY
                        '%d %b %Y'   # DD Mon YYYY
                    ]
                    
                    start_date = None
                    end_date = None
                    
                    # Traiter le début
                    for fmt in date_formats:
                        try:
                            start_date = datetime.strptime(str(start), fmt)
                            break
                        except (ValueError, TypeError):
                            continue
                    
                    # Traiter la fin
                    if isinstance(end, str) and end.lower() in ['aujourd\'hui', 'present', 'actuel', 'maintenant', 'current', 'now', 'today', 'présent']:
                        end_date = datetime.now()
                    else:
                        for fmt in date_formats:
                            try:
                                end_date = datetime.strptime(str(end), fmt)
                                break
                            except (ValueError, TypeError):
                                continue
                    
                    if start_date and end_date:
                        # Calculer la différence en mois
                        months = ((end_date.year - start_date.year) * 12) + (end_date.month - start_date.month)
                        
                        # Ajouter les jours proportionnellement
                        if end_date.day > start_date.day:
                            months += (end_date.day - start_date.day) / 30
                        
                        # Si c'est un stage, vérifier la durée typique
                        if is_internship and months > 6:
                            months = 6
                            
                        experience_years.append(months / 12)
                except Exception as e:
                    print(f"Erreur lors du calcul de la durée: {str(e)}")
        except Exception as e:
            print(f"Erreur lors du traitement de l'entrée {start}-{end}: {str(e)}")
    
    # 4. Calculer l'expérience totale
    total_experience = sum(experience_years) if experience_years else 0
    
    # 5. Vérifier s'il s'agit d'un profil junior/étudiant
    is_junior = re.search(r'stage|stagiaire|intern|internship|junior|entry[- ]level|débutant|étudiant', resume_text.lower()) is not None
    
    # 6. Détection de l'âge pour vérification de cohérence
    age = None
    
    # Recherche directe de l'âge
    age_pattern = r'(?:âge|age)\s*:?\s*(\d{1,2})|(\d{1,2})\s*(?:ans|years)'
    age_match = re.search(age_pattern, resume_text.lower())
    if age_match:
        age_group = age_match.group(1) or age_match.group(2)
        if age_group:
            age = int(age_group)
    
    # Recherche par date de naissance
    birth_year = None
    
    # Format JJ/MM/AAAA
    birth_date_pattern = r'(?:né\(e\)|né|naissance|birth|born)(?:\s+(?:le|on|:))?\s*:?\s*(\d{1,2})[\/\.-](\d{1,2})[\/\.-](\d{4}|\d{2})'
    birth_match = re.search(birth_date_pattern, resume_text.lower())
    if birth_match:
        try:
            year = int(birth_match.group(3))
            if year < 100:  # Année sur 2 chiffres
                year = 1900 + year if year > 50 else 2000 + year
            birth_year = year
            age = datetime.now().year - birth_year
        except (ValueError, IndexError):
            pass
    
    # Format "Date de naissance : 02 Mars 2005"
    birth_date_text_pattern = r'(?:date\s+de\s+naissance|birth\s+date).*?(\d{1,2})[\s]*(?:janvier|février|mars|avril|mai|juin|juillet|août|septembre|octobre|novembre|décembre|january|february|march|april|may|june|july|august|september|october|november|december)[\s]*(\d{4})'
    birth_text_match = re.search(birth_date_text_pattern, resume_text.lower())
    if birth_text_match:
        try:
            birth_year = int(birth_text_match.group(2))
            age = datetime.now().year - birth_year
        except (ValueError, IndexError):
            pass
    
    # 7. Vérifier la cohérence avec l'âge
    if age:
        # Estimation de l'expérience maximale possible (âge - 18)
        max_possible_experience = max(0, age - 18)
        
        # Si l'expérience détectée dépasse l'expérience maximale possible, ajuster
        if total_experience > max_possible_experience:
            total_experience = max_possible_experience
    
    # 8. Vérifier la cohérence avec la date de diplôme la plus récente
    graduation_year = None
    recent_graduation = re.search(r'(?:diplôme|graduate|graduation|obtention).*?(20\d{2})', resume_text.lower())
    if recent_graduation:
        try:
            graduation_year = int(recent_graduation.group(1))
            years_since_graduation = datetime.now().year - graduation_year
            
            if years_since_graduation < 3 and total_experience > years_since_graduation:
                total_experience = years_since_graduation
        except:
            pass
    
    # 9. Arrondir à une décimale pour plus de précision
    final_experience = round(total_experience, 1)
    
    # 10. Ajustement spécial pour les CV avec uniquement des stages
    if is_junior and re.search(r'stagiaire', resume_text.lower()) and not re.search(r'développeur|ingénieur|consultant|analyste|manager|directeur', resume_text.lower()):
        if final_experience > 1.0:
            stage_months = 0
            for exp in experience_years:
                # Pour chaque expérience, si c'est probablement un stage, limiter à 6 mois max
                if exp <= 0.5:  # 6 mois ou moins
                    stage_months += exp * 12
                else:
                    # Si plus de 6 mois, vérifier le contexte
                    if is_junior:
                        stage_months += min(exp, 0.5) * 12
                    else:
                        stage_months += exp * 12
            
            final_experience = round(stage_months / 12, 1)
    
    return final_experience

def extract_candidate_skills_no_llm(resume_text):
    """
    Extrait les compétences d'un CV sans utiliser de LLM.
    Version de secours basée sur l'analyse du texte en temps réel.
    
    Args:
        resume_text: Texte du CV
            
    Returns:
        Dictionnaire contenant les compétences du candidat
    """
    skills = []
    education = []
    
    # Listes de technologies et compétences courantes à rechercher
    common_techs = ["Python", "JavaScript", "TypeScript", "Java", "C#", "PHP", 
                    "Ruby", "Go", "React", "Angular", "Vue", "Node.js", "Django", 
                    "Flask", "Spring", "Express", "Laravel", "Ruby on Rails", 
                    "PostgreSQL", "MySQL", "MongoDB", "SQL Server", "Oracle", 
                    "Flutter", "Dart", "Swift", "Kotlin"]
    
    common_skills = ["API REST", "GraphQL", "Git", "Docker", "Kubernetes", "CI/CD", 
                     "AWS", "Azure", "GCP", "DevOps", "Agile", "Scrum", "TDD", 
                     "Machine Learning", "Deep Learning", "Data Science", "Big Data", 
                     "Cloud Computing", "Microservices", "Mobile Development",
                     "UML", "RUP", "Tests unitaires", "UI/UX"]
    
    # Recherche basique de mots-clés et détermination du niveau
    technical_skills = []
    
    # Analyser chaque ligne pour les compétences et niveaux d'expertise
    lines = resume_text.split('\n')
    for line in lines:
        # Chercher les compétences techniques
        for tech in common_techs:
            if tech.lower() in line.lower():
                level = "intermédiaire"  # niveau par défaut
                
                # Détecter le niveau basé sur le contexte
                line_lower = line.lower()
                if "expert" in line_lower or "avancée" in line_lower or "confirmé" in line_lower:
                    level = "expert"
                elif "avancé" in line_lower or "forte" in line_lower or "solide" in line_lower:
                    level = "avancé"
                elif "débutant" in line_lower or "notions" in line_lower or "base" in line_lower:
                    level = "débutant"
                
                skills.append({"name": tech, "level": level})
                if tech not in technical_skills:
                    technical_skills.append(tech)
        
        # Chercher les autres compétences
        for skill in common_skills:
            if skill.lower() in line.lower():
                level = "intermédiaire"  # niveau par défaut
                
                # Détecter le niveau basé sur le contexte
                line_lower = line.lower()
                if "expert" in line_lower or "avancée" in line_lower or "confirmé" in line_lower:
                    level = "expert"
                elif "avancé" in line_lower or "forte" in line_lower or "solide" in line_lower:
                    level = "avancé"
                elif "débutant" in line_lower or "notions" in line_lower or "base" in line_lower:
                    level = "débutant"
                
                skills.append({"name": skill, "level": level})
                if skill not in technical_skills:
                    technical_skills.append(skill)
    
    # Détection plus avancée de l'éducation
    education_data = []
    education_keywords = [
        "Licence", "Master", "Doctorat", "DUT", "BTS", "Ingénieur", "Université", "École", "Formation", "Diplôme",
        "Bachelor", "Master's", "PhD", "Degree", "BSc", "MSc", "MBA", "University", "College", "Education", "Certificate"

    ]
    year_pattern = r'20\d{2}|19\d{2}'  # Années entre 1900 et 2099
    
    for i, line in enumerate(lines):
        for keyword in education_keywords:
            if keyword.lower() in line.lower():
                # Obtenir plus d'informations en analysant la ligne et les lignes suivantes et précédentes
                context_start = max(0, i - 2)
                context_end = min(len(lines), i + 3)
                context = ' '.join(lines[context_start:context_end])
                
                # Extraire l'année (si disponible)
                year_match = re.search(year_pattern, context)
                year = year_match.group(0) if year_match else "N/A"
                
                # Extraire l'institution (essai simple)
                institution = "Non spécifiée"
                institution_keywords = ["Université", "École", "Institute", "College", "Faculté"]
                for inst_kw in institution_keywords:
                    if inst_kw.lower() in context.lower():
                        # Extraire le nom complet de l'institution
                        inst_pattern = f"{inst_kw}[^,;\n.]*"
                        inst_match = re.search(inst_pattern, context, re.IGNORECASE)
                        if inst_match:
                            institution = inst_match.group(0).strip()
                            break
                
                # Déterminer la pertinence basée sur les mots-clés dans la formation
                relevance = "moyenne"
                if any(tech.lower() in context.lower() for tech in technical_skills):
                    relevance = "haute"
                
                # Éviter les doublons
                if not any(edu["degree"] == line.strip() for edu in education_data):
                    education_data.append({
                        "degree": line.strip(),
                        "institution": institution,
                        "year": year,
                        "relevance": relevance
                    })
                break
    
    # Détection de l'expérience professionnelle
    experience_years = 0
    roles = []
    relevant_experience = []
    
    # Recherche des années d'expérience
    experience_patterns = [
        r'(\d+)\s*(?:an(?:s|née(?:s)?)?\s+d\'?(?:expérience))',
        r'expérience\s+(?:de|:)?\s*(\d+)\s*an(?:s|née(?:s)?)?',
        r'(?:depuis|pendant)\s+(\d+)\s*an(?:s|née(?:s)?)?',
        
        r'(\d+)\s*(?:year(?:s)?\s+(?:of\s+)?experience)',
        r'experience\s+(?:of|:)?\s*(\d+)\s*year(?:s)?',
        r'(?:for|over)\s+(\d+)\s*year(?:s)?'
    ]
    
    for pattern in experience_patterns:
        matches = re.finditer(pattern, resume_text.lower())
        for match in matches:
            try:
                years = int(match.group(1))
                experience_years = max(experience_years, years)
            except:
                pass
    
    # Recherche de postes et entreprises
    company_patterns = [
        r'(?:chez|à|pour|au sein de)\s+([A-Z][A-Za-z0-9\s&\.\-]+)',
        r'([A-Z][A-Za-z0-9\s&\.\-]+)(?:\s*,|\s+SA|\s+SAS|\s+SARL|\s+Inc\.)',
        
        r'(?:at|for|with)\s+([A-Z][A-Za-z0-9\s&\.\-]+)',
        r'([A-Z][A-Za-z0-9\s&\.\-]+)(?:\s*,|\s+Inc|\s+Ltd|\s+LLC|\s+Corp|\s+Limited)'
    ]
    
    role_patterns = [
        r'([A-Z][a-zéèêë]+\s+(?:de\s+)?[Dd]éveloppeur\w*)',
        r'([Dd]éveloppeur\w*\s+[A-Za-z]+)',
        r'([A-Z][a-zéèêë]+\s+(?:d\'|en\s+)?[Ii]ngénieur\w*)',
        r'([Ii]ngénieur\w*\s+[A-Za-z]+)',
        r'([Cc]hef\s+de\s+projet)',
        r'([Cc]onsultant\w*\s+[A-Za-z]+)',
        r'([Aa]rchitecte\s+[A-Za-z]+)',
        r'([Aa]nalyst\w*\s+[A-Za-z]+)',
        r'([Dd]ata\s+[Ss]cientist)',
        r'([Tt]ech\s*(?:nical)?\s*[Ll]ead)',
        r'([Dd]ev[Oo]ps)',
        r'([Ff]ull\s*[Ss]tack)',
        r'([Ff]ront\s*[Ee]nd)',
        r'([Bb]ack\s*[Ee]nd)',
        
        r'([Ss]oftware\s+[Ee]ngineer\w*)',
        r'([Pp]roduct\s+[Mm]anager\w*)',
        r'([Pp]roject\s+[Mm]anager\w*)',
        r'([Ss]enior\s+[A-Za-z\s]+(?:Developer|Engineer|Architect|Consultant))',
        r'([Jj]unior\s+[A-Za-z\s]+(?:Developer|Engineer|Analyst))',
        r'([Ll]ead\s+[A-Za-z\s]+(?:Developer|Engineer|Architect))',
        r'([Tt]echnical\s+[A-Za-z\s]+)',
        r'([Pp]rogrammer\w*)',
        r'([Ss]ystems?\s+[Aa]dministrator\w*)',
        r'([Nn]etwork\s+[Ee]ngineer\w*)',
        r'([Ss]ecurity\s+[A-Za-z\s]+)',
        r'([Qq]uality\s+[Aa]ssurance\w*)',
        r'([Tt]ester\w*)',
        r'([Uu][Xx]/[Uu][Ii]\s+[Dd]esigner\w*)'
    ]
    
    # Extraire les rôles
    for pattern in role_patterns:
        matches = re.finditer(pattern, resume_text)
        for match in matches:
            role = match.group(1).strip()
            if role not in roles:
                roles.append(role)
    
    # Rechercher les expériences professionnelles (rôle + entreprise + durée)
    position_blocks = []
    current_block = {}
    
    # Analyser ligne par ligne pour trouver des blocs d'expérience
    for i, line in enumerate(lines):
        # Chercher les rôles/titres de poste
        for pattern in role_patterns:
            role_match = re.search(pattern, line)
            if role_match and (not current_block or 'position' not in current_block):
                current_block = {'position': role_match.group(1).strip()}
                # Chercher l'entreprise sur la même ligne ou les lignes suivantes
                company_context = ' '.join(lines[i:min(i+3, len(lines))])
                for company_pattern in company_patterns:
                    company_match = re.search(company_pattern, company_context)
                    if company_match:
                        current_block['company'] = company_match.group(1).strip()
                        break
                
                # Chercher la durée
                duration_match = re.search(r'(\d+)\s*(?:an(?:s|née(?:s)?)?)', company_context)
                if duration_match:
                    current_block['duration'] = f"{duration_match.group(1)} ans"
                else:
                    # Chercher des dates pour calculer la durée
                    years_match = re.findall(r'(20\d{2}|19\d{2})', company_context)
                    if len(years_match) >= 2:
                        years = sorted([int(y) for y in years_match])
                        duration = years[-1] - years[0]
                        current_block['duration'] = f"{duration} ans"
                    else:
                        current_block['duration'] = "Durée non spécifiée"
                
                # Ajouter des points clés
                highlights = []
                highlight_context = ' '.join(lines[i:min(i+10, len(lines))])
                # Rechercher des réalisations marquées par des puces ou numéros
                highlight_matches = re.findall(r'(?:•|\-|\*|–|—|\d+[\.\)])\s+([^\n\r•\-\*]+)', highlight_context)
                for highlight in highlight_matches[:3]:  # Prendre les 3 premiers points
                    if len(highlight.strip()) > 10:  # Ignorer les points trop courts
                        highlights.append(highlight.strip())
                
                # S'il n'y a pas de points formatés, essayer d'extraire des phrases clés
                if not highlights:
                    sentences = re.split(r'[.!?]+', highlight_context)
                    for sentence in sentences[:3]:
                        if any(kw in sentence.lower() for kw in ['développ', 'mis en place', 'créé', 'réalis', 'implémen', 'conçu']):
                            if len(sentence.strip()) > 15:
                                highlights.append(sentence.strip())
                
                # Ajouter des valeurs par défaut si nécessaire
                if 'company' not in current_block:
                    current_block['company'] = "Entreprise non spécifiée"
                if 'duration' not in current_block:
                    current_block['duration'] = "Durée non spécifiée"
                if not highlights:
                    highlights = ["Expérience professionnelle pertinente"]
                
                current_block['highlights'] = highlights
                position_blocks.append(current_block)
                current_block = {}
    
    # Si des expériences ont été trouvées, les utiliser
    if position_blocks:
        relevant_experience = position_blocks
    # Sinon, créer des entrées basées sur les rôles détectés
    elif roles:
        for role in roles:
            relevant_experience.append({
                "position": role,
                "company": "Détectée à partir du CV",
                "duration": f"{max(1, experience_years // len(roles))} ans",
                "highlights": ["Expérience professionnelle pertinente"]
            })
    
    # Détection des soft skills
    soft_skills = []
    common_soft_skills = ["Communication", "Travail d'équipe", "Autonomie", 
                          "Résolution de problèmes", "Adaptabilité", "Créativité",
                          "Leadership", "Gestion du temps", "Organisation", 
                          "Rigueur", "Polyvalence", "Esprit d'analyse", 
                          "Sens des responsabilités", "Capacité d'apprentissage"]
    
    # Rechercher les soft skills dans le texte complet
    for skill in common_soft_skills:
        if skill.lower() in resume_text.lower():
            soft_skills.append(skill)
    
    # Inférer des soft skills supplémentaires basés sur le contexte
    if any(kw in resume_text.lower() for kw in ["équipe", "collabor", "groupe"]):
        if "Travail d'équipe" not in soft_skills:
            soft_skills.append("Travail d'équipe")
    
    if any(kw in resume_text.lower() for kw in ["autonome", "indépend", "seul"]):
        if "Autonomie" not in soft_skills:
            soft_skills.append("Autonomie")
    
    if any(kw in resume_text.lower() for kw in ["complex", "problèm", "solution", "résoudre"]):
        if "Résolution de problèmes" not in soft_skills:
            soft_skills.append("Résolution de problèmes")
    
    # Valeurs minimales par défaut uniquement si rien n'est détecté
    if not technical_skills:
        technical_skills = []
    
    if not soft_skills:
        soft_skills = []
    
    if not education_data:
        education_data = []
    
    if not relevant_experience:
        relevant_experience = []
    
    # Créer un résumé basé sur les données extraites
    resume_summary = ""
    
    # Ajouter les rôles et l'expérience au résumé
    if roles:
        resume_summary += f"Professionnel avec expérience en tant que {', '.join(roles[:2])}. "
    elif experience_years > 0:
        resume_summary += f"Professionnel avec {experience_years} ans d'expérience. "
    
    # Ajouter les compétences techniques au résumé
    if technical_skills:
        resume_summary += f"Compétences en {', '.join(technical_skills[:5])}. "
    
    # Ajouter la formation si disponible
    if education_data:
        degrees = [edu['degree'] for edu in education_data]
        resume_summary += f"Formation: {', '.join(degrees[:2])}."
    
    # Calculer un score d'adéquation basé sur des facteurs objectifs
    # Plus de poids aux compétences techniques et à l'expérience
    fit_score = min(7.0, 3.0 + 
                   min(2.0, len(technical_skills) * 0.2) + 
                   min(1.0, len(soft_skills) * 0.1) + 
                   min(1.0, experience_years * 0.1))
    
    # Identifier les forces basées sur les compétences et l'expérience
    strengths = []
    
    # Ajouter les compétences techniques comme forces
    if technical_skills:
        top_skills = technical_skills[:3]
        strengths.append(f"Compétences techniques en {', '.join(top_skills)}")
    
    # Ajouter l'expérience comme force si significative
    if experience_years > 2:
        strengths.append(f"Expérience professionnelle de {experience_years} ans")
    elif relevant_experience:
        strengths.append("Expérience professionnelle pertinente")
    
    # Ajouter les soft skills comme forces
    if soft_skills:
        top_soft_skills = soft_skills[:2]
        strengths.append(f"Compétences en {', '.join(top_soft_skills)}")
    
    # Ajouter la formation comme force si pertinente
    if education_data and any(edu["relevance"] == "haute" for edu in education_data):
        strengths.append("Formation académique pertinente")
    
    # Identifier les lacunes potentielles
    gaps = []
    
    # Manque d'expérience
    if experience_years < 2 and not any("stage" in exp.get("position", "").lower() for exp in relevant_experience):
        gaps.append("Expérience professionnelle limitée")
    
    # Manque de formation spécifique si pertinent
    if not education_data:
        gaps.append("Formation non spécifiée dans le CV")
    
    # Si très peu de compétences détectées
    if len(technical_skills) < 3:
        gaps.append("Gamme limitée de compétences techniques détectées")
    
    # Générer des questions pertinentes
    recommended_questions = []
    
    # Question sur l'expérience professionnelle
    if relevant_experience:
        most_recent_role = relevant_experience[0]["position"]
        recommended_questions.append({
            "question": f"Pouvez-vous me détailler votre expérience en tant que {most_recent_role}?",
            "rationale": "Pour évaluer l'expérience professionnelle récente"
        })
    
    # Question sur les compétences techniques
    if technical_skills:
        main_skill = technical_skills[0]
        recommended_questions.append({
            "question": f"Pouvez-vous me parler d'un projet où vous avez utilisé {main_skill}?",
            "rationale": f"Pour évaluer l'expertise pratique en {main_skill}"
        })
    
    # Question sur la résolution de problèmes
    recommended_questions.append({
        "question": "Décrivez un défi technique complexe que vous avez rencontré et comment vous l'avez résolu.",
        "rationale": "Pour évaluer les capacités de résolution de problèmes et l'approche technique"
    })
    
    # Retourner les résultats avec toutes les données extraites
    return {
        "resume_summary": resume_summary,
        "technical_skills": technical_skills,
        "soft_skills": soft_skills,
        "relevant_experience": relevant_experience,
        "education": education_data,
        "fit_score": fit_score,
        "fit_justification": "Score basé sur l'analyse des compétences techniques, soft skills et expérience professionnelle",
        "strengths": strengths,
        "gaps": gaps,
        "recommended_questions": recommended_questions,
        "keywords": technical_skills + soft_skills,
        "_generated_by": "fallback_method"  # Marqueur pour indiquer que c'est la méthode de secours
    }

def extract_job_requirements_no_llm(job_text):
    """
    Extrait les exigences d'une offre d'emploi sans utiliser de LLM.
    Version de secours basée sur l'analyse du texte en temps réel.
    
    Args:
        job_text: Texte de l'offre d'emploi
            
    Returns:
        Dictionnaire contenant les exigences du poste
    """
    # Listes de technologies et compétences courantes à rechercher
    common_tech_stacks = ["Python", "JavaScript", "TypeScript", "Java", "C#", "PHP", 
                         "Ruby", "Go", "React", "Angular", "Vue", "Node.js", "Django", 
                         "Flask", "Spring", "Express", "Laravel", "Ruby on Rails", 
                         "PostgreSQL", "MySQL", "MongoDB", "SQL Server", "Oracle", 
                         "Flutter", "Dart", "Swift", "Kotlin", ".NET","HTML/CSS"]
    
    common_skills = ["API REST", "GraphQL", "Git", "Docker", "Kubernetes", "CI/CD", 
                    "AWS", "Azure", "GCP", "DevOps", "Agile", "Scrum", "TDD", 
                    "Machine Learning", "Deep Learning", "Data Science", "Big Data", 
                    "Cloud Computing", "Microservices", "Mobile Development",
                    "UML", "RUP", "Tests unitaires", "UX/UI"]
    
    # Initialiser les collections pour les exigences
    tech_stacks = []
    technical_skills = []
    soft_skills = []
    certifications = []
    
    # Diviser le texte en sections pour une analyse plus précise
    lines = job_text.split('\n')
    
    # Rechercher les sections courantes dans les descriptions de poste
    sections = {
        "requirements": [],
        "qualifications": [],
        "skills": [],
        "profile": [],
        "experience": [],
        "education": [],
        "about": []
    }
    
    current_section = "requirements"  # Section par défaut
    
    # Identifier les sections dans le texte
    for line in lines:
        line_lower = line.lower().strip()

        # Détecter les en-têtes de section (français et anglais)
        if any(header in line_lower for header in ["requis", "exigé", "demandé", "required", "requirements", "necessity"]):
            current_section = "requirements"
            continue
        elif any(header in line_lower for header in ["qualifications", "profil", "profile", "candidate"]):
            current_section = "qualifications"
            continue
        elif any(header in line_lower for header in ["compétences", "skills", "abilities", "competencies", "what you'll need"]):
            current_section = "skills"
            continue
        elif any(header in line_lower for header in ["expérience", "experience"]):
            current_section = "experience"
            continue
        elif any(header in line_lower for header in ["formation", "études", "education", "diplôme", "degree", "academic"]):
            current_section = "education"
            continue
        elif any(header in line_lower for header in ["entreprise", "société", "about us", "qui sommes-nous", "company", "organization"]):
            current_section = "about"
            continue
        
        # Ajouter la ligne à la section courante
        sections[current_section].append(line)
    
    # Fusionner toutes les sections pertinentes pour la recherche de compétences
    relevant_sections = ' '.join([
        ' '.join(sections["requirements"]),
        ' '.join(sections["qualifications"]),
        ' '.join(sections["skills"]),
        ' '.join(sections["profile"]),
        ' '.join(sections["experience"])
    ])
    
    # Rechercher les technologies et compétences
    for tech in common_tech_stacks:
        if re.search(r'\b' + re.escape(tech.lower()) + r'\b', relevant_sections.lower()):
            tech_stacks.append(tech)
    
    for skill in common_skills:
        if re.search(r'\b' + re.escape(skill.lower()) + r'\b', relevant_sections.lower()):
            technical_skills.append(skill)
    
    # Détecter les exigences de piles technologiques supplémentaires
    # (expressions comme "React/Angular/Vue" ou "React ou Angular ou Vue")
    stack_groups = re.findall(r'([A-Za-z]+(?:\s*[/,]\s*|\s+ou\s+)[A-Za-z]+(?:\s*[/,]\s*|\s+ou\s+)[A-Za-z]+)', relevant_sections)
    for group in stack_groups:
        # Séparer les technologies dans le groupe
        group_techs = re.split(r'\s*[/,]\s*|\s+ou\s+', group)
        for tech in group_techs:
            if tech and tech not in tech_stacks and any(t.lower() == tech.lower() for t in common_tech_stacks):
                matching_tech = next(t for t in common_tech_stacks if t.lower() == tech.lower())
                tech_stacks.append(matching_tech)
    
    # Rechercher les années d'expérience requises
    experience_patterns = [
        r'(\d+)[+\-]?\s*(?:an(?:s|née(?:s)?)?\s+d\'?(?:expérience))',
        r'expérience\s+(?:de|:)?\s*(\d+)[+\-]?\s*an(?:s|née(?:s)?)?',
        r'(?:minimum|au moins)\s+(\d+)\s*an(?:s|née(?:s)?)?'
    ]
    
    experience_years = []
    for pattern in experience_patterns:
        matches = re.finditer(pattern, job_text.lower())
        for match in matches:
            try:
                years = int(match.group(1))
                experience_years.append(years)
            except:
                pass
    
    # Déterminer le niveau d'expérience
    experience_level = "Junior"
    if experience_years:
        max_years = max(experience_years)
        if max_years >= 5:
            experience_level = "5+ ans"
        elif max_years >= 3:
            experience_level = "3-5 ans"
        else:
            experience_level = "0-2 ans"
    else:
        # Détecter le niveau d'expérience à partir des mots-clés
        if any(kw in job_text.lower() for kw in ["senior", "expérimenté", "confirmé", "5 ans", "5+"]):
            experience_level = "5+ ans"
        elif any(kw in job_text.lower() for kw in ["intermédiaire", "3 ans", "3-5", "3 à 5"]):
            experience_level = "3-5 ans"
        elif any(kw in job_text.lower() for kw in ["junior", "débutant", "entry level", "graduate"]):
            experience_level = "0-2 ans"
    
    # Rechercher les soft skills
    common_soft_skills = ["Communication", "Travail d'équipe", "Autonomie", 
                          "Résolution de problèmes", "Adaptabilité", "Créativité",
                          "Leadership", "Gestion du temps", "Organisation", 
                          "Rigueur", "Polyvalence", "Esprit d'analyse", 
                          "Sens des responsabilités", "Capacité d'apprentissage"]
    
    for skill in common_soft_skills:
        if skill.lower() in job_text.lower():
            soft_skills.append(skill)
    
    # Inférer des soft skills supplémentaires basés sur le contexte
    if any(kw in job_text.lower() for kw in ["équipe", "collabor", "groupe"]):
        if "Travail d'équipe" not in soft_skills:
            soft_skills.append("Travail d'équipe")
    
    if any(kw in job_text.lower() for kw in ["autonome", "indépend", "seul"]):
        if "Autonomie" not in soft_skills:
            soft_skills.append("Autonomie")
    
    if any(kw in job_text.lower() for kw in ["complex", "problèm", "solution", "résoudre"]):
        if "Résolution de problèmes" not in soft_skills:
            soft_skills.append("Résolution de problèmes")
    
    # Rechercher les certifications
    common_certifications = ["AWS Certified", "Microsoft Certified", "Google Cloud", 
                            "Scrum Master", "PMP", "CISSP", "CISA", "ITIL", "TOGAF"]
    
    for cert in common_certifications:
        if cert.lower() in job_text.lower():
            certifications.append(cert)
    
    # Extraire des certifications supplémentaires selon un pattern
    cert_patterns = [
        r'[Cc]ertifi(?:é|cation)\s+([A-Za-z0-9\s]+)',
        r'[Cc]ertification\s+(?:en|:)\s+([A-Za-z0-9\s]+)'
    ]
    
    for pattern in cert_patterns:
        matches = re.finditer(pattern, job_text)
        for match in matches:
            cert = match.group(1).strip()
            if cert and cert not in certifications:
                certifications.append(cert)
    
    return {
        "tech_stacks": tech_stacks,
        "technical_skills": technical_skills,
        "soft_skills": soft_skills,
        "experience_level": experience_level,
        "certifications": certifications,
        "_generated_by": "fallback_method"  # Marqueur pour indiquer que c'est la méthode de secours
    }

def match_resume_to_job_description_no_llm(resume_text, job_description):
    """
    Compare un CV à une description de poste sans utiliser de LLM.
    Version de secours basée sur l'analyse approfondie en temps réel.
    
    Args:
        resume_text: Le texte du CV
        job_description: La description du poste
        
    Returns:
        dict: Résultats de correspondance
    """
    # Extraire les compétences du CV et les exigences du poste
    resume_data = extract_candidate_skills_no_llm(resume_text)
    job_data = extract_job_requirements_no_llm(job_description)
    
    # Normaliser les compétences pour la comparaison
    def normalize_skills(skills):
        return [skill.lower() for skill in skills]
    
    # Compétences techniques du CV (normalisées)
    resume_skills = set(normalize_skills(resume_data["technical_skills"]))
    
    # Compétences techniques requises pour le poste (normalisées)
    job_skills = set(normalize_skills(job_data["technical_skills"] + job_data["tech_stacks"]))
    
    # Compétences correspondantes (originales, non normalisées)
    matching_skills = [
        skill for skill in (resume_data["technical_skills"])
        if skill.lower() in job_skills
    ]
    
    # Compétences manquantes (originales, non normalisées)
    missing_skills = [
        skill for skill in (job_data["technical_skills"] + job_data["tech_stacks"])
        if skill.lower() not in resume_skills
    ]
    
    # Analyse des soft skills
    resume_soft_skills = set(normalize_skills(resume_data["soft_skills"]))
    job_soft_skills = set(normalize_skills(job_data["soft_skills"]))
    matching_soft_skills = [
        skill for skill in resume_data["soft_skills"]
        if skill.lower() in job_soft_skills
    ]
    
    # Analyse de l'expérience
    experience_match = 0
    
    # Calculer les années d'expérience du candidat
    candidate_experience = 0
    # Essayer d'extraire les années d'expérience des données du CV
    if "experience" in resume_data and "years" in resume_data["experience"]:
        candidate_experience = resume_data["experience"]["years"]
    else:
        # Estimer à partir des expériences professionnelles
        for exp in resume_data["relevant_experience"]:
            duration = exp.get("duration", "")
            years_match = re.search(r'(\d+)', duration)
            if years_match:
                candidate_experience += int(years_match.group(1))
    
    # Déterminer l'expérience requise pour le poste
    required_experience = 0
    if job_data["experience_level"] == "0-2 ans":
        required_experience = 1
    elif job_data["experience_level"] == "3-5 ans":
        required_experience = 4
    elif job_data["experience_level"] == "5+ ans":
        required_experience = 6
    
    # Calculer la correspondance d'expérience
    if candidate_experience >= required_experience:
        experience_match = 1.0  # Correspondance parfaite ou supérieure
    else:
        experience_match = max(0, candidate_experience / required_experience)
    
    # Analyse de l'adéquation des compétences
    # Calculer l'importance relative des compétences du poste
    total_job_skills = len(job_skills)
    if total_job_skills == 0:
        skills_match_score = 0.5  # Score moyen par défaut
    else:
        # Calculer le nombre de compétences correspondantes
        skills_match_count = len([skill for skill in resume_skills if skill in job_skills])
        skills_match_score = skills_match_count / total_job_skills
    
    # Calculer la correspondance des soft skills
    total_job_soft_skills = len(job_soft_skills)
    if total_job_soft_skills == 0:
        soft_skills_match_score = 0.5  # Score moyen par défaut
    else:
        # Calculer le nombre de soft skills correspondantes
        soft_skills_match_count = len([skill for skill in resume_soft_skills if skill in job_soft_skills])
        soft_skills_match_score = soft_skills_match_count / total_job_soft_skills
    
    # Calculer le score de correspondance global pondéré
    # Plus de poids aux compétences techniques, puis à l'expérience, puis aux soft skills
    match_score = (
        (skills_match_score * 0.6) +  # 60% pour les compétences techniques
        (experience_match * 0.3) +    # 30% pour l'expérience
        (soft_skills_match_score * 0.1)  # 10% pour les soft skills
    ) * 100  # Convertir en pourcentage
    
    # Arrondir le score
    match_score = round(match_score)
    
    # Identifier les points forts du candidat
    key_strengths = []
    
    # Ajouter les compétences correspondantes comme points forts
    if matching_skills:
        key_strengths.append(f"Maîtrise des technologies requises: {', '.join(matching_skills[:3])}")
    
    # Ajouter l'expérience si adéquate
    if candidate_experience >= required_experience:
        key_strengths.append(f"Expérience adéquate: {candidate_experience} ans (requis: {required_experience} ans)")
    
    # Ajouter les soft skills correspondantes
    if matching_soft_skills:
        key_strengths.append(f"Compétences interpersonnelles requises: {', '.join(matching_soft_skills[:3])}")
    
    # Ajouter la formation si pertinente
    if resume_data["education"] and any(edu["relevance"] == "haute" for edu in resume_data["education"]):
        key_strengths.append("Formation académique pertinente pour le poste")
    
    # S'assurer d'avoir au moins 2 points forts
    if len(key_strengths) < 2:
        # Chercher d'autres points forts potentiels
        if resume_data["technical_skills"]:
            key_strengths.append(f"Compétences techniques variées incluant {', '.join(resume_data['technical_skills'][:3])}")
        
        if resume_data["relevant_experience"]:
            key_strengths.append("Expérience professionnelle dans le domaine")
    
    # Déterminer l'évaluation globale
    assessment = ""
    if match_score >= 80:
        assessment = "Excellente correspondance - Le candidat possède la plupart des compétences et l'expérience requises pour le poste."
    elif match_score >= 60:
        assessment = "Bonne correspondance - Le candidat possède plusieurs compétences clés requises mais pourrait nécessiter une formation dans certains domaines."
    elif match_score >= 40:
        assessment = "Correspondance moyenne - Le candidat possède quelques compétences pertinentes mais manque d'expérience ou de compétences clés."
    else:
        assessment = "Correspondance faible - Le profil du candidat diffère significativement des exigences du poste."
    
    # Ajouter des informations contextuelles à l'évaluation
    if len(matching_skills) > 0:
        match_ratio = len(matching_skills) / max(1, len(job_data["technical_skills"] + job_data["tech_stacks"]))
        if match_ratio > 0.7:
            assessment += f" Le candidat maîtrise {len(matching_skills)} des {len(job_data['technical_skills'] + job_data['tech_stacks'])} compétences techniques requises."
        else:
            assessment += f" Le candidat maîtrise seulement {len(matching_skills)} des {len(job_data['technical_skills'] + job_data['tech_stacks'])} compétences techniques requises."
    
    if candidate_experience < required_experience:
        assessment += f" Le niveau d'expérience ({candidate_experience} ans) est inférieur aux {required_experience} ans requis."
    
    # Assembler le résultat
    return {
        "match_score": match_score,
        "matching_skills": matching_skills,
        "missing_skills": missing_skills,
        "key_strengths": key_strengths,
        "overall_assessment": assessment,
        "_generated_by": "fallback_method"  # Marqueur pour indiquer que c'est la méthode de secours
    }

def generate_interview_questions_from_resume_no_llm(resume_text, job_role, num_questions=5):
    """
    Génère des questions d'entretien personnalisées sans utiliser de LLM.
    Version de secours basée sur l'analyse réelle du CV.
    
    Args:
        resume_text: Le texte du CV
        job_role: Le poste pour lequel le candidat postule
        num_questions: Nombre de questions à générer
        
    Returns:
        list: Liste de questions d'entretien
    """
    # Extraire les compétences et expériences du CV
    skills_data = extract_candidate_skills_no_llm(resume_text)
    
    # Récupérer les informations clés
    technical_skills = skills_data.get("technical_skills", [])
    soft_skills = skills_data.get("soft_skills", [])
    experiences = skills_data.get("relevant_experience", [])
    education = skills_data.get("education", [])
    
    # Collection de questions
    questions = []
    
    # 1. Question sur l'expérience professionnelle la plus récente ou la plus pertinente
    if experiences:
        latest_experience = experiences[0]
        position = latest_experience.get("position", "votre dernier poste")
        company = latest_experience.get("company", "votre précédente entreprise")
        
        questions.append({
            "question": f"Pouvez-vous me décrire vos principales responsabilités et réalisations en tant que {position} chez {company}?",
            "type": "expérience",
            "skill_to_assess": "Expérience professionnelle",
            "difficulty": "facile",
            "expected_answer_elements": ["Responsabilités clés", "Projets significatifs", "Réalisations mesurables", "Compétences techniques utilisées"]
        })
    else:
        # Question générique sur l'expérience si aucune expérience spécifique n'est détectée
        questions.append({
            "question": f"Pouvez-vous me parler de votre parcours professionnel et de son adéquation avec le poste de {job_role}?",
            "type": "expérience",
            "skill_to_assess": "Parcours professionnel",
            "difficulty": "facile",
            "expected_answer_elements": ["Expériences pertinentes", "Progression de carrière", "Motivation pour le poste actuel"]
        })
    
    # 2. Questions sur les compétences techniques spécifiques
    if technical_skills:
        # Sélectionner jusqu'à 2 compétences techniques pour les questions
        for skill in technical_skills[:min(2, len(technical_skills))]:
            # Questions adaptées à différentes compétences
            if skill.lower() in ["python", "java", "javascript", "c#", "php", "ruby", "go"]:
                questions.append({
                    "question": f"Décrivez un projet complexe où vous avez utilisé {skill}. Quels défis techniques avez-vous rencontrés et comment les avez-vous surmontés?",
                    "type": "technique",
                    "skill_to_assess": f"Expertise en {skill}",
                    "difficulty": "moyenne",
                    "expected_answer_elements": ["Description du projet", "Défis techniques", "Solutions implémentées", "Résultats obtenus"]
                })
            elif skill.lower() in ["react", "angular", "vue", "flutter"]:
                questions.append({
                    "question": f"Comment structurez-vous une application {skill} pour qu'elle soit maintenable et évolutive? Pouvez-vous donner un exemple concret?",
                    "type": "technique",
                    "skill_to_assess": f"Architecture {skill}",
                    "difficulty": "difficile",
                    "expected_answer_elements": ["Structure de l'application", "Gestion de l'état", "Bonnes pratiques", "Exemple concret"]
                })
            elif skill.lower() in ["sql", "mysql", "postgresql", "mongodb", "oracle"]:
                questions.append({
                    "question": f"Comment optimisez-vous les performances des requêtes dans {skill}? Donnez un exemple de requête que vous avez optimisée.",
                    "type": "technique",
                    "skill_to_assess": f"Optimisation {skill}",
                    "difficulty": "difficile",
                    "expected_answer_elements": ["Techniques d'indexation", "Optimisation de requêtes", "Analyse de performance", "Exemple concret"]
                })
            else:
                questions.append({
                    "question": f"Quelle est votre expérience avec {skill} et comment l'avez-vous appliquée dans un contexte professionnel?",
                    "type": "technique",
                    "skill_to_assess": f"Expertise en {skill}",
                    "difficulty": "moyenne",
                    "expected_answer_elements": ["Niveau d'expertise", "Applications pratiques", "Contexte d'utilisation", "Résultats obtenus"]
                })
    
    # 3. Question sur la résolution de problèmes (toujours pertinente)
    questions.append({
        "question": "Décrivez une situation où vous avez dû résoudre un problème technique particulièrement complexe. Quelle était votre approche et quel a été le résultat?",
        "type": "situation",
        "skill_to_assess": "Résolution de problèmes",
        "difficulty": "moyenne",
        "expected_answer_elements": ["Analyse du problème", "Méthodologie", "Solutions envisagées", "Solution retenue", "Résultat final"]
    })
    
    # 4. Question sur le travail d'équipe ou la communication (si mentionné dans les soft skills)
    if "Travail d'équipe" in soft_skills or "Communication" in soft_skills:
        questions.append({
            "question": "Pouvez-vous me donner un exemple de situation où vos compétences en communication ou en travail d'équipe ont été essentielles pour la réussite d'un projet?",
            "type": "comportementale",
            "skill_to_assess": "Communication et travail d'équipe",
            "difficulty": "moyenne",
            "expected_answer_elements": ["Contexte du projet", "Défis de communication", "Actions entreprises", "Impact sur le projet"]
        })
    
    # 5. Question sur la gestion du temps ou des priorités
    questions.append({
        "question": "Comment gérez-vous vos priorités lorsque vous avez plusieurs tâches ou projets simultanés avec des délais serrés?",
        "type": "comportementale",
        "skill_to_assess": "Gestion du temps et des priorités",
        "difficulty": "moyenne",
        "expected_answer_elements": ["Méthodes d'organisation", "Outils utilisés", "Définition des priorités", "Gestion du stress"]
    })
    
    # 6. Question sur l'apprentissage continu (importante pour le développement)
    questions.append({
        "question": "Comment vous tenez-vous à jour avec les dernières technologies et tendances dans votre domaine? Donnez un exemple d'une nouvelle compétence que vous avez récemment acquise.",
        "type": "comportementale",
        "skill_to_assess": "Apprentissage continu",
        "difficulty": "facile",
        "expected_answer_elements": ["Sources d'information", "Méthodes d'apprentissage", "Exemple concret", "Application pratique"]
    })
    
    # 7. Question spécifique au poste
    questions.append({
        "question": f"Selon vous, quelles sont les qualités et compétences les plus importantes pour réussir en tant que {job_role}?",
        "type": "comportementale",
        "skill_to_assess": "Compréhension du rôle",
        "difficulty": "facile",
        "expected_answer_elements": ["Compétences techniques pertinentes", "Soft skills nécessaires", "Compréhension des responsabilités", "Vision du poste"]
    })
    
    # 8. Question sur un projet ou une réalisation dont le candidat est fier
    questions.append({
        "question": "Quel projet ou quelle réalisation professionnelle vous a procuré le plus de satisfaction et pourquoi?",
        "type": "comportementale",
        "skill_to_assess": "Motivation et fierté professionnelle",
        "difficulty": "facile",
        "expected_answer_elements": ["Description du projet", "Contribution personnelle", "Défis surmontés", "Impact et résultats", "Apprentissages"]
    })
    
    # 9. Question sur les défis et l'adaptabilité
    questions.append({
        "question": "Face à un changement imprévu dans les spécifications ou les priorités d'un projet, comment réagissez-vous et adaptez-vous votre travail?",
        "type": "situation",
        "skill_to_assess": "Adaptabilité et gestion du changement",
        "difficulty": "moyenne",
        "expected_answer_elements": ["Réaction initiale", "Communication avec l'équipe", "Ajustement du plan", "Flexibilité"]
    })
    
    # 10. Question sur les aspirations professionnelles
    questions.append({
        "question": "Où vous voyez-vous professionnellement dans les 3 à 5 prochaines années et comment ce poste s'inscrit-il dans votre plan de carrière?",
        "type": "comportementale",
        "skill_to_assess": "Ambition et alignement avec l'entreprise",
        "difficulty": "facile",
        "expected_answer_elements": ["Objectifs professionnels", "Plan de développement", "Alignement avec le poste actuel", "Vision à long terme"]
    })
    
    # Limiter au nombre de questions demandé
    return questions[:num_questions]

def analyze_resume(file, job_role=None):
   """
   Analyse un CV téléchargé et retourne des informations structurées
   
   Args:
       file: Le fichier CV téléchargé
       job_role (str, optional): Le poste pour lequel le candidat postule
       
   Returns:
       dict: Résultats d'analyse structurés
   """
   try:
       # Vérifier si le fichier est valide
       if not file or not allowed_file(file.filename):
           return {
               "error": "Format de fichier non pris en charge. Veuillez télécharger un fichier PDF, DOCX ou TXT."
           }
       
       # Extraire le texte du CV
       resume_text = extract_text_from_file(file)
       
       language = detect_language(resume_text)

       
       # Prétraiter le texte
       processed_text = preprocess_resume_text(resume_text)
       
       # Extraire les informations de contact
       contact_info = extract_contact_info(processed_text)
       
       # Si aucun poste n'est spécifié, utiliser un poste générique
       if job_role is None:
           job_role = "un poste dans votre entreprise" if language == 'fr' else "a position in your company"
       
       
       # Tenter d'analyser le CV avec le LLM
       try:
           prompt = RESUME_ANALYSIS_PROMPT.format(
               job_role=job_role,
               resume_text=processed_text[:8000]  # Limiter la taille pour éviter de dépasser le contexte du LLM
           )
           
           llm_response = get_llm_response(prompt)
           
           try:
               # Tenter de parser la réponse JSON
               analysis = json.loads(llm_response)
               
               # Ajouter les informations de contact et métadonnées
               analysis['contact_info'] = contact_info
               analysis['metadata'] = {
                   'filename': secure_filename(file.filename),
                   'analysis_timestamp': datetime.now().isoformat(),
                   'job_role': job_role,
                   'analysis_method': 'llm'
               }
               
               return analysis
               
           except json.JSONDecodeError:
               # Si la réponse n'est pas au format JSON valide, utiliser la méthode de secours
               print(f"Erreur de décodage JSON dans la réponse LLM, utilisation de la méthode de secours")
               fallback_analysis = extract_candidate_skills_no_llm(processed_text)
               fallback_analysis['contact_info'] = contact_info
               fallback_analysis['metadata'] = {
                   'filename': secure_filename(file.filename),
                   'analysis_timestamp': datetime.now().isoformat(),
                   'job_role': job_role,
                   'analysis_method': 'fallback',
                   'error': 'JSON invalide dans la réponse LLM'
               }
               
               return fallback_analysis
               
       except Exception as llm_error:
           # En cas d'erreur avec le LLM, utiliser la méthode de secours
           print(f"Erreur LLM: {str(llm_error)}, utilisation de la méthode de secours")
           fallback_analysis = extract_candidate_skills_no_llm(processed_text)
           fallback_analysis['contact_info'] = contact_info
           fallback_analysis['metadata'] = {
               'filename': secure_filename(file.filename),
               'analysis_timestamp': datetime.now().isoformat(),
               'job_role': job_role,
               'analysis_method': 'fallback',
               'error': f'Erreur LLM: {str(llm_error)}'
           }
           
           return fallback_analysis
           
   except Exception as e:
       # Capturer toute autre erreur et essayer la méthode de secours
       try:
           print(f"Erreur générale: {str(e)}, tentative d'utilisation de la méthode de secours")
           # Si une erreur se produit mais que nous avons déjà le texte du CV
           if 'processed_text' in locals():
               fallback_analysis = extract_candidate_skills_no_llm(processed_text)
               fallback_analysis['contact_info'] = extract_contact_info(processed_text) if 'contact_info' not in locals() else contact_info
               fallback_analysis['metadata'] = {
                   'filename': secure_filename(file.filename),
                   'analysis_timestamp': datetime.now().isoformat(),
                   'job_role': job_role,
                   'analysis_method': 'fallback',
                   'error': f'Erreur générale: {str(e)}'
               }
               return fallback_analysis
       except:
           pass
           
       # Si même la méthode de secours échoue ou si nous n'avons pas le texte du CV
       return {
           "error": f"Erreur lors de l'analyse du CV: {str(e)}",
           "analysis_method": "failed"
       }

def match_resume_to_job_description(resume_text, job_description, detailed=True):
   """
   Compare un CV à une description de poste et évalue la correspondance
   
   Args:
       resume_text (str): Le texte du CV
       job_description (str): La description du poste
       detailed (bool): Si True, retourne une analyse détaillée
       
   Returns:
       dict: Résultats de correspondance
   """
   try:
       prompt = f"""
       Tu es un expert en recrutement chargé d'évaluer l'adéquation entre un CV et une description de poste.

       CV:
       ```
       {resume_text[:4000]}  # Limiter pour ne pas dépasser le contexte
       ```

       Description du poste:
       ```
       {job_description[:2000]}  # Limiter pour ne pas dépasser le contexte
       ```

       {"Analyse de façon détaillée" if detailed else "Analyse brièvement"} l'adéquation entre ce CV et cette description de poste. 
       
       Fournis ta réponse au format JSON avec la structure suivante:
       {{
         "match_score": 85,  // Score de correspondance sur 100
         "matching_skills": ["compétence1", "compétence2", ...],  // Compétences présentes dans le CV qui correspondent au poste
         "missing_skills": ["compétence1", "compétence2", ...],  // Compétences importantes pour le poste mais absentes du CV
         "key_strengths": ["point1", "point2", ...],  // Points forts du candidat pour ce poste
         "overall_assessment": "Évaluation globale de l'adéquation"
       }}
       """
       
       try:
           llm_response = get_llm_response(prompt)
           match_analysis = json.loads(llm_response)
           match_analysis['analysis_method'] = 'llm'
           return match_analysis
       except Exception as llm_error:
           # En cas d'erreur avec le LLM, utiliser la méthode de secours
           print(f"Erreur LLM dans match_resume_to_job_description: {str(llm_error)}, utilisation de la méthode de secours")
           fallback_analysis = match_resume_to_job_description_no_llm(resume_text, job_description)
           fallback_analysis['analysis_method'] = 'fallback'
           fallback_analysis['error'] = f'Erreur LLM: {str(llm_error)}'
           return fallback_analysis
   except Exception as e:
       # En cas d'erreur générale, utiliser la méthode de secours
       try:
           print(f"Erreur générale dans match_resume_to_job_description: {str(e)}, utilisation de la méthode de secours")
           fallback_analysis = match_resume_to_job_description_no_llm(resume_text, job_description)
           fallback_analysis['analysis_method'] = 'fallback'
           fallback_analysis['error'] = f'Erreur générale: {str(e)}'
           return fallback_analysis
       except:
           return {
               "error": f"Erreur lors de l'analyse de correspondance: {str(e)}",
               "match_score": 0,
               "analysis_method": "failed"
           }

def generate_interview_questions_from_resume(resume_text, job_role, num_questions=5):
   """
   Génère des questions d'entretien personnalisées basées sur le CV
   
   Args:
       resume_text (str): Le texte du CV
       job_role (str): Le poste pour lequel le candidat postule
       num_questions (int): Nombre de questions à générer
       
   Returns:
       list: Liste de questions d'entretien
   """
   try:
       prompt = f"""
       Tu es un recruteur expérimenté spécialisé dans les entretiens pour le poste de {job_role}.
       Après avoir analysé le CV suivant, génère {num_questions} questions d'entretien pertinentes 
       qui permettront d'évaluer les compétences et l'expérience du candidat.

       CV:
       ```
       {resume_text[:5000]}  # Limiter pour ne pas dépasser le contexte
       ```

       Les questions doivent:
       1. Être spécifiques au parcours et aux compétences mentionnés dans le CV
       2. Inclure un mélange de questions techniques et comportementales
       3. Permettre d'évaluer l'adéquation au poste de {job_role}
       4. Être formulées de manière ouverte pour encourager des réponses détaillées

       Fournis ta réponse au format JSON avec la structure suivante:
       [
         {{
           "question": "Question d'entretien complète",
           "type": "technique/comportementale/situation/expérience",
           "skill_to_assess": "Compétence évaluée",
           "difficulty": "facile/moyenne/difficile",
           "expected_answer_elements": ["élément1", "élément2", ...] // Points clés à rechercher dans la réponse
         }},
         ...
       ]
       """
       
       try:
           llm_response = get_llm_response(prompt)
           questions = json.loads(llm_response)
           # Ajouter des métadonnées d'analyse
           for question in questions:
               question['generated_by'] = 'llm'
           return questions
       except Exception as llm_error:
           # En cas d'erreur avec le LLM, utiliser la méthode de secours
           print(f"Erreur LLM dans generate_interview_questions: {str(llm_error)}, utilisation de la méthode de secours")
           fallback_questions = generate_interview_questions_from_resume_no_llm(resume_text, job_role, num_questions)
           # Ajouter des métadonnées d'analyse
           for question in fallback_questions:
               question['generated_by'] = 'fallback'
               question['error'] = f'Erreur LLM: {str(llm_error)}'
           return fallback_questions
   except Exception as e:
       # En cas d'erreur générale, utiliser la méthode de secours
       try:
           print(f"Erreur générale dans generate_interview_questions: {str(e)}, utilisation de la méthode de secours")
           fallback_questions = generate_interview_questions_from_resume_no_llm(resume_text, job_role, num_questions)
           # Ajouter des métadonnées d'analyse
           for question in fallback_questions:
               question['generated_by'] = 'fallback'
               question['error'] = f'Erreur générale: {str(e)}'
           return fallback_questions
       except:
           # Dernière ressource : questions génériques
           return [
               {
                   "question": f"Parlez-moi de votre expérience en tant que {job_role}.",
                   "type": "expérience",
                   "skill_to_assess": "Expérience générale",
                   "difficulty": "facile",
                   "expected_answer_elements": ["Parcours professionnel", "Réalisations"],
                   "generated_by": "default",
                   "error": f"Erreur lors de la génération des questions: {str(e)}"
               },
               {
                   "question": "Décrivez un projet difficile sur lequel vous avez travaillé récemment.",
                   "type": "situation",
                   "skill_to_assess": "Résolution de problèmes",
                   "difficulty": "moyenne",
                   "expected_answer_elements": ["Analyse de la situation", "Approche méthodique", "Résultats"],
                   "generated_by": "default",
                   "error": f"Erreur lors de la génération des questions : {str(e)}"
               }
           ]

def get_additional_context_for_improvement(improvement_target, job_description=None, experience_years=None):
    """
    Génère le contexte additionnel à inclure dans le prompt d'amélioration basé sur la cible
    
    Args:
        improvement_target (str): La cible d'amélioration ("job_offer", "american", "canadian", "linkedin")
        job_description (str, optional): Description du poste si improvement_target est "job_offer"
        experience_years (int, optional): Années d'expérience détectées
        
    Returns:
        str: Contexte additionnel pour le prompt
    """
    experience_context = ""
    if experience_years is not None and experience_years > 0:
        experience_context = f"""
        Le CV indique {experience_years} ans d'expérience professionnelle. 
        Assure-toi que cette information soit clairement mise en valeur dans le résumé et dans l'en-tête du CV.
        """
    
    if improvement_target == "job_offer" and job_description:
        return f"""
        Voici l'offre d'emploi pour laquelle le CV doit être optimisé:
        ```
        {job_description}
        ```
        
        Analyse la correspondance entre le CV et cette offre d'emploi.
        Identifie les compétences et expériences qui manquent ou devraient être mises en avant.
        {experience_context}
        Suggère des modifications pour augmenter les chances du candidat d'être retenu par les systèmes ATS 
        (Applicant Tracking System) et les recruteurs.
        """
    
    elif improvement_target == "american":
        return f"""
        Ce CV doit être adapté au style américain. Les CV américains:
        - Mettent l'accent sur les résultats quantifiables et les accomplissements
        - Utilisent des verbes d'action forts
        - Sont concis (1-2 pages maximum)
        - Sont orientés vers les compétences et l'expérience pertinentes pour le poste visé
        - N'incluent généralement pas de photo, d'âge, de statut marital ou d'informations personnelles
        - Commencent souvent par un résumé professionnel ou un objectif de carrière ciblé
        - Utilisent le format antichronologique pour l'expérience professionnelle et la formation
        {experience_context}
        """
    
    elif improvement_target == "canadian":
        return f"""
        Ce CV doit être adapté au style canadien. Les CV canadiens:
        - Sont similaires aux CV américains, mais peuvent inclure plus d'informations personnelles
        - Mettent l'accent sur les compétences transférables
        - Incluent souvent une section sur les langues parlées (surtout pour les postes au Québec)
        - Peuvent mentionner le statut d'immigration si pertinent
        - Sont généralement limités à 2 pages maximum
        - Préfèrent un ton formel et professionnel
        - Peuvent inclure des expériences de bénévolat
        {experience_context}
        """
    
    elif improvement_target == "linkedin":
        return f"""
        Ce CV doit être adapté au format LinkedIn. Un profil LinkedIn efficace:
        - Utilise un titre/headline accrocheur qui combine poste actuel et proposition de valeur
        - Commence par une section "À propos" rédigée à la première personne, conversationnelle mais professionnelle
        - Quantifie les réalisations dans chaque expérience professionnelle
        - Utilise des mots-clés pertinents pour le secteur et recherchés par les recruteurs
        - Inclut des compétences que les connexions peuvent valider
        - Limite chaque section à l'essentiel pour être facilement scannable
        - Peut inclure des médias, articles, projets pour illustrer les réalisations
        {experience_context}
        """
    
    else:
        return f"""
        Analyse ce CV et suggère des améliorations générales pour le rendre plus efficace et percutant.
        Concentre-toi sur la structure, le contenu, les formulations et la présentation.
        {experience_context}
        """

def improve_cv_no_llm(resume_text, improvement_target, job_description=None):
    """
    Version de secours pour améliorer un CV sans utiliser de LLM.
    
    Args:
        resume_text (str): Le texte du CV
        improvement_target (str): La cible d'amélioration ("job_offer", "american", "canadian", "linkedin")
        job_description (str, optional): Description du poste si improvement_target est "job_offer"
        
    Returns:
        dict: Analyse et suggestions d'amélioration
    """
    # Initialiser les collections pour l'analyse
    strengths = []
    weaknesses = []
    missing_elements = []
    format_issues = []
    content_improvements = []
    structure_improvements = []
    keyword_optimization = []
    ats_compatibility_tips = []
    
    # Analyse de base du CV
    lines = resume_text.split('\n')
    text_length = len(resume_text)
    word_count = len(resume_text.split())
    
    # Détection des sections standard dans le CV
    has_summary = any(re.search(r'(résumé|sommaire|profil|summary|about|à propos|profile|professional summary)', line, re.IGNORECASE) for line in lines[:15])
    has_skills = any(re.search(r'(compétences|skills|expertise|technologies|core competencies)', line, re.IGNORECASE) for line in lines)
    has_experience = any(re.search(r'(expérience|experience|parcours|carrière|work history|employment)', line, re.IGNORECASE) for line in lines)
    has_education = any(re.search(r'(formation|education|études|diplôme|academic|degree)', line, re.IGNORECASE) for line in lines)
    has_languages = any(re.search(r'(langues|languages|language skills)', line, re.IGNORECASE) for line in lines)
    has_certifications = any(re.search(r'(certifications|certificats|diplômes|qualifications)', line, re.IGNORECASE) for line in lines)
    
    # Recherche de verbes d'action
    action_verbs = ["développé", "dirigé", "géré", "créé", "implémenté", "conçu", "coordonné", 
                    "analyzed", "managed", "developed", "implemented", "coordinated", "created",
                    "achieved", "improved", "increased", "decreased", "reduced"]
    
    action_verbs_count = sum(1 for verb in action_verbs if re.search(r'\b' + verb + r'\b', resume_text, re.IGNORECASE))
    
    # Recherche de quantifications
    quantification_pattern = r'\b\d+%|\d+ [a-zA-Z]+|\$\d+|\d+ \$'
    quantification_matches = re.findall(quantification_pattern, resume_text)
    has_quantifications = len(quantification_matches) > 0
        
    # Évaluation basique de la structure du CV
    if text_length < 1500:
        format_issues.append("Le CV est trop court et manque probablement de détails importants")
    elif text_length > 8000:
        format_issues.append("Le CV est trop long et devrait être condensé pour plus d'impact")
    
    # Points forts identifiés
    if has_summary:
        strengths.append("Inclut un résumé professionnel/sommaire en début de CV")
    else:
        missing_elements.append("Résumé professionnel ou sommaire de compétences")
    
    if has_skills:
        strengths.append("Présente une section dédiée aux compétences")
    else:
        missing_elements.append("Section de compétences techniques et/ou interpersonnelles")
    
    if has_experience:
        strengths.append("Détaille l'expérience professionnelle")
    else:
        missing_elements.append("Expérience professionnelle détaillée")
    
    if has_education:
        strengths.append("Inclut les informations sur la formation/éducation")
    else:
        missing_elements.append("Section sur la formation académique")
    
   
    
    # Points faibles identifiés
    if action_verbs_count < 5:
        weaknesses.append("Utilisation insuffisante de verbes d'action percutants")
        content_improvements.append({
            "section": "Expérience professionnelle",
            "current_content": "Descriptions de postes sans verbes d'action",
            "suggested_improvement": "Commencer chaque point par un verbe d'action fort et percutant",
            "rationale": "Les verbes d'action rendent le CV plus dynamique et mettent en valeur vos contributions"
        })
        keyword_optimization.extend(["développé", "dirigé", "géré", "implémenté", "optimisé"])
    
    if not has_quantifications:
        weaknesses.append("Manque de résultats quantifiés et mesurables")
        content_improvements.append({
            "section": "Expérience professionnelle",
            "current_content": "Descriptions de responsabilités sans mesures concrètes",
            "suggested_improvement": "Ajouter des métriques et des pourcentages pour quantifier vos réalisations",
            "rationale": "Les chiffres attirent l'attention et démontrent votre impact de façon concrète"
        })
    
    final_experience_years = extract_experience_years(resume_text)
    
    # Vérifier si l'expérience est clairement mentionnée dans le CV
    if final_experience_years > 0:
        # Vérifier si les années d'expérience sont clairement indiquées dans le résumé
        exp_in_summary = re.search(r'\b' + str(final_experience_years) + r'\s*(?:an(?:s|née(?:s)?)|year(?:s)?)', resume_text[:500], re.IGNORECASE)
        if not exp_in_summary:
            weaknesses.append(f"L'expérience totale de {final_experience_years} ans n'est pas clairement mise en avant en début de CV")
            content_improvements.append({
                "section": "Résumé professionnel",
                "current_content": "Résumé sans mention claire de l'expérience totale",
                "suggested_improvement": f"Mentionner explicitement '{final_experience_years} ans d'expérience en [domaine]' dans le résumé",
                "rationale": "L'expérience professionnelle est un atout clé qui doit être immédiatement visible"
            })
    else:
        if has_experience:
            weaknesses.append("Années d'expérience non clairement indiquées")
            content_improvements.append({
                "section": "Résumé professionnel",
                "current_content": "Absence de mention d'années d'expérience",
                "suggested_improvement": "Inclure une phrase indiquant clairement votre expérience totale dans votre domaine",
                "rationale": "Les recruteurs recherchent rapidement cette information clé"
            })
    
    # Recommandations spécifiques selon la cible d'amélioration
    if improvement_target == "job_offer" and job_description:
        # [Code existant pour job_offer]
        job_data = extract_job_requirements_no_llm(job_description)
        required_skills = job_data.get("technical_skills", []) + job_data.get("tech_stacks", [])
        required_soft_skills = job_data.get("soft_skills", [])
        job_exp_level = job_data.get("experience_level", "")
        
        # Extraire les compétences du candidat
        resume_data = extract_candidate_skills_no_llm(resume_text)
        candidate_skills = resume_data.get("technical_skills", [])
        candidate_soft_skills = resume_data.get("soft_skills", [])
        
        # Identifier les compétences manquantes
        missing_technical_skills = [skill for skill in required_skills if skill.lower() not in [s.lower() for s in candidate_skills]]
        missing_soft_skills = [skill for skill in required_soft_skills if skill.lower() not in [s.lower() for s in candidate_soft_skills]]
        
        if missing_technical_skills:
            weaknesses.append("Certaines compétences techniques requises ne sont pas mentionnées")
            keyword_optimization.extend(missing_technical_skills)
        
        if missing_soft_skills:
            weaknesses.append("Certaines compétences non-techniques requises ne sont pas mises en valeur")
            keyword_optimization.extend(missing_soft_skills)
        
        # Vérifier l'adéquation de l'expérience
        job_years_required = 0
        if "0-2 ans" in job_exp_level:
            job_years_required = 1
        elif "3-5 ans" in job_exp_level:
            job_years_required = 3
        elif "5+ ans" in job_exp_level:
            job_years_required = 5
        
        if job_years_required > 0:
            if final_experience_years < job_years_required:
                weaknesses.append(f"L'expérience de {final_experience_years} ans est inférieure aux {job_years_required} ans requis pour le poste")
                content_improvements.append({
                    "section": "Expérience professionnelle",
                    "current_content": f"Expérience totale de {final_experience_years} ans",
                    "suggested_improvement": "Mettre en valeur tous les projets et réalisations pertinents pour compenser le manque d'années d'expérience",
                    "rationale": "Démontrer que vous pouvez apporter de la valeur malgré une expérience formelle plus courte"
                })
            else:
                strengths.append(f"L'expérience de {final_experience_years} ans correspond ou dépasse les {job_years_required} ans requis pour le poste")
                content_improvements.append({
                    "section": "Résumé professionnel",
                    "current_content": "Mention générique de l'expérience",
                    "suggested_improvement": f"Souligner vos {final_experience_years} ans d'expérience, dépassant les {job_years_required} ans requis",
                    "rationale": "Mettre en avant un avantage compétitif clé par rapport aux exigences du poste"
                })
        
        # Conseils pour l'optimisation ATS avec expérience
        ats_compatibility_tips = [
            "Inclure les mots-clés exacts de l'offre d'emploi dans votre CV",
            "Utiliser à la fois la version complète et l'acronyme des technologies (ex: 'JavaScript (JS)')",
            "Adapter le titre/poste recherché pour qu'il corresponde exactement au titre dans l'offre",
            "Éviter les tableaux, en-têtes/pieds de page et formats complexes qui peuvent poser problème aux ATS",
            "Utiliser un format de fichier standard comme .docx ou .pdf (texte sélectionnable)",
            f"Mentionner clairement vos {final_experience_years} ans d'expérience directement dans le résumé et l'en-tête"
        ]
        
        structure_improvements.append({
            "recommendation": "Réorganiser le CV pour mettre en avant les compétences et expériences les plus pertinentes pour ce poste",
            "importance": "élevée",
            "details": "Placer les compétences et expériences correspondant le mieux à l'offre en début de section"
        })
    
    elif improvement_target in ["american", "canadian", "linkedin"]:
        template = CV_TEMPLATES.get(improvement_target, {})
        expected_sections = template.get("sections", [])
        format_guidelines = template.get("format_guidelines", {})
        
        # Vérifier les sections attendues pour ce format
        for section in expected_sections:
            section_pattern = "|".join([word.lower() for word in re.split(r'[_\s]', section)])
            if not any(re.search(section_pattern, line, re.IGNORECASE) for line in lines):
                missing_elements.append(f"Section '{section}'")
        
        # Conseils spécifiques au format
        if improvement_target == "american":
            if not has_quantifications:
                structure_improvements.append({
                    "recommendation": "Ajouter des résultats quantifiables à chaque expérience professionnelle",
                    "importance": "élevée",
                    "details": "Les CV américains mettent l'accent sur les résultats mesurables"
                })
            personal_info_patterns = [
                r'âge|age',
                r'date de naissance',
                r'marié|marrié',
                r'photo'
            ]
            if any(re.search(pattern, resume_text, re.IGNORECASE) for pattern in personal_info_patterns):
                format_issues.append("Les informations personnelles comme l'âge, le statut marital ou les photos ne sont pas recommandées dans les CV américains")    
        elif improvement_target == "canadian":
            if not has_languages:
                missing_elements.append("Section sur les langues parlées (important au Canada, surtout au Québec)")
            
            structure_improvements.append({
                "recommendation": "Ajouter une section sur les langues maîtrisées",
                "importance": "moyenne",
                "details": "Préciser le niveau de maîtrise pour chaque langue (ex: Français - natif, Anglais - courant)"
            })
            
        elif improvement_target == "linkedin":
            structure_improvements.append({
                "recommendation": "Rédiger la section 'À propos' à la première personne",
                "importance": "élevée",
                "details": "Sur LinkedIn, un ton plus conversationnel et personnel est apprécié"
            })
            
            if not has_summary or not any(line.strip() and len(line.split()) > 20 for line in lines[:15]):
                content_improvements.append({
                    "section": "À propos / Summary",
                    "current_content": "Résumé trop court ou inexistant",
                    "suggested_improvement": "Rédiger un résumé percutant de 3-5 phrases à la première personne",
                    "rationale": "La section 'À propos' est cruciale sur LinkedIn pour capter l'attention des recruteurs"
                })
    
    
    # Créer un résumé professionnel amélioré générique avec années d'expérience
    career_level = "professionnel expérimenté"
    experience_description = ""
    
    if final_experience_years > 0:
        experience_description = f"avec {final_experience_years} ans d'expérience"
        if final_experience_years < 3:
            career_level = f"jeune professionnel {experience_description}"
        elif final_experience_years < 7:
            career_level = f"professionnel {experience_description}"
        else:
            career_level = f"professionnel expérimenté {experience_description}"
    else:
        career_level = "professionnel"
    
    # Extraire les compétences principales 
    skills_data = extract_candidate_skills_no_llm(resume_text)
    main_skills = skills_data.get("technical_skills", [])[:5]
    soft_skills = skills_data.get("soft_skills", [])[:3]
    
    # Générer quelques réalisations génériques basées sur les compétences
    achievements = []
    for skill in main_skills[:2]:
        achievements.append(f"expertise en {skill}")
    
    # Construire le résumé amélioré
    skills_part = f"spécialisé en {', '.join(main_skills)}" if main_skills else "aux compétences diversifiées"
    soft_part = f" et possédant d'excellentes aptitudes en {', '.join(soft_skills)}" if soft_skills else ""
    achievements_part = f", avec une {' et une '.join(achievements)}" if achievements else ""
    
    improved_summary = f"{career_level} {skills_part}{soft_part}{achievements_part}. Orienté résultats et solutions, capable de relever des défis complexes et d'atteindre les objectifs fixés."
    
    # Assembler le résultat
    return {
        "analysis": {
            "strengths": strengths,
            "weaknesses": weaknesses,
            "missing_elements": missing_elements,
            "format_issues": format_issues
        },
        "recommendations": {
            "content_improvements": content_improvements,
            "structure_improvements": structure_improvements,
            "keyword_optimization": keyword_optimization,
            "ats_compatibility_tips": ats_compatibility_tips if ats_compatibility_tips and len(ats_compatibility_tips) > 0 else []
        },
        "experience_years": final_experience_years,
        "improved_summary": improved_summary,
        "_generated_by": "fallback_method"  # Marqueur pour indiquer que c'est la méthode de secours
    }

def format_cv_content_no_llm(resume_text, improvement_target, job_description=None, cv_analysis=None):
    """
    Reformate le contenu d'un CV sans utiliser de LLM
    
    Args:
        resume_text (str): Le texte du CV
        improvement_target (str): La cible d'amélioration
        job_description (str, optional): Description du poste
        cv_analysis (dict, optional): Analyse précédente du CV
        
    Returns:
        str: Contenu du CV reformaté
    """
    # Utiliser l'analyse si fournie, sinon en générer une nouvelle
    if not cv_analysis:
        cv_analysis = improve_cv_no_llm(resume_text, improvement_target, job_description)
    
    # Extraire les informations clés du CV original
    lines = resume_text.split('\n')
    
    # Extraire les informations de contact
    contact_info = extract_contact_info(resume_text)
    
    # Extraire les compétences et expériences
    skills_data = extract_candidate_skills_no_llm(resume_text)
    
    # Sélectionner le modèle approprié
    template = CV_TEMPLATES.get(improvement_target, CV_TEMPLATES.get("american"))
    
    # Créer la structure du nouveau CV
    formatted_cv = []
    
    # En-tête avec informations de contact
    name_line = next((line for line in lines[:5] if len(line.strip()) > 0 and len(line.strip()) < 50), "")
    if name_line:
        formatted_cv.append(name_line.strip().upper())
        formatted_cv.append("")
    
    contact_line = []
    if contact_info.get('email'):
        contact_line.append(contact_info['email'])
    if contact_info.get('phone'):
        contact_line.append(contact_info['phone'])
    if contact_info.get('linkedin'):
        contact_line.append(contact_info['linkedin'])
    
    if contact_line:
        formatted_cv.append(" | ".join(contact_line))
        formatted_cv.append("")
    
    # Résumé professionnel
    improved_summary = cv_analysis.get('improved_summary', '')
    if improved_summary:
        formatted_cv.append(template['sections'][0])  # Premier section (PROFESSIONAL SUMMARY/ABOUT)
        formatted_cv.append("------------------------------")
        formatted_cv.append(improved_summary)
        formatted_cv.append("")
    
    # Compétences
    if 'SKILLS' in template['sections'] or 'COMPÉTENCES' in template['sections']:
        section_title = 'SKILLS' if improvement_target in ['american', 'linkedin'] else 'COMPÉTENCES'
        formatted_cv.append(section_title)
        formatted_cv.append("------------------------------")

        # Ajouter les compétences techniques
        technical_skills = skills_data.get('technical_skills', [])
        if technical_skills:
            if improvement_target == 'linkedin':
                # Format LinkedIn: liste de compétences
                formatted_cv.append("Technical: " + ", ".join(technical_skills))
            else:
                # Format standard: sections
                formatted_cv.append("Technical Skills:" if improvement_target in ['american', 'linkedin'] else "Compétences techniques:")
                skills_chunks = [technical_skills[i:i+5] for i in range(0, len(technical_skills), 5)]
                for chunk in skills_chunks:
                    formatted_cv.append("• " + ", ".join(chunk))

        # Ajouter les soft skills
        soft_skills = skills_data.get('soft_skills', [])
        if soft_skills:
            if improvement_target == 'linkedin':
                formatted_cv.append("Soft Skills: " + ", ".join(soft_skills))
            else:
                formatted_cv.append("Soft Skills:" if improvement_target in ['american', 'linkedin'] else "Compétences personnelles:")
                formatted_cv.append("• " + ", ".join(soft_skills))
        
        formatted_cv.append("")
    
    # Expérience professionnelle
    if 'EXPERIENCE' in template['sections'] or 'EXPÉRIENCE PROFESSIONNELLE' in template['sections']:
        section_title = 'EXPERIENCE' if improvement_target in ['american', 'linkedin'] else 'EXPÉRIENCE PROFESSIONNELLE'
        formatted_cv.append(section_title)
        formatted_cv.append("------------------------------")
        
        # Ajouter les expériences
        experiences = skills_data.get('relevant_experience', [])
        if experiences:
            for exp in experiences:
                position = exp.get('position', '')
                company = exp.get('company', '')
                duration = exp.get('duration', '')
                
                # Titre du poste et entreprise
                if position and company:
                    if improvement_target == 'linkedin':
                        formatted_cv.append(f"{position}")
                        formatted_cv.append(f"{company} | {duration}")
                    else:
                        formatted_cv.append(f"{position}, {company} ({duration})")
                
                # Points clés de l'expérience
                highlights = exp.get('highlights', [])
                for highlight in highlights:
                    # Améliorer les points avec des verbes d'action
                    first_word = highlight.strip().split(' ')[0].lower()
                    action_verbs = ["développé", "dirigé", "géré", "créé", "implémenté", "conçu", "coordonné"]
                    
                    if first_word not in [verb.lower() for verb in action_verbs]:
                        # Ajouter un verbe d'action si aucun n'est présent
                        import random
                        verb = random.choice(action_verbs)
                        highlight = f"{verb} {highlight[0].lower()}{highlight[1:]}"
                    
                    formatted_cv.append(f"• {highlight}")
                
                formatted_cv.append("")
        else:
            # Expérience générique si aucune n'est détectée
            formatted_cv.append("[Insérez vos expériences ici en suivant ce format:]")
            formatted_cv.append("Titre du poste, Entreprise (Durée)")
            formatted_cv.append("• Réalisation importante avec résultats quantifiables")
            formatted_cv.append("• Deuxième accomplissement majeur")
            formatted_cv.append("")
    
    # Formation
    if 'EDUCATION' in template['sections'] or 'FORMATION' in template['sections']:
        section_title = 'EDUCATION' if improvement_target in ['american', 'linkedin'] else 'FORMATION'
        formatted_cv.append(section_title)
        formatted_cv.append("------------------------------")
        
        # Ajouter la formation
        education = skills_data.get('education', [])
        if education:
            for edu in education:
                degree = edu.get('degree', '')
                institution = edu.get('institution', '')
                year = edu.get('year', '')
                
                if degree and institution:
                    formatted_cv.append(f"{degree}, {institution} ({year})")
                    formatted_cv.append("")
        else:
            # Formation générique si aucune n'est détectée
            formatted_cv.append("[Insérez votre formation ici en suivant ce format:]")
            formatted_cv.append("Diplôme, Institution (Année)")
            formatted_cv.append("")
    
    # Sections spécifiques selon le format
    if improvement_target == 'canadian' and 'LANGUAGES' in template['sections']:
        formatted_cv.append("LANGUAGES")
        formatted_cv.append("------------------------------")
        formatted_cv.append("[Insérez vos langues et niveaux ici]")
        formatted_cv.append("")
    
    if 'CERTIFICATIONS' in template['sections']:
        formatted_cv.append("CERTIFICATIONS")
        formatted_cv.append("------------------------------")
        formatted_cv.append("[Insérez vos certifications ici]")
        formatted_cv.append("")
    
    # Conseils d'optimisation ATS
    if improvement_target == 'job_offer':
        formatted_cv.append("")
        formatted_cv.append("NOTES SUR L'OPTIMISATION ATS (à retirer avant utilisation)")
        formatted_cv.append("------------------------------")
        for tip in cv_analysis.get('recommendations', {}).get('ats_compatibility_tips', []):
            formatted_cv.append(f"• {tip}")
    
    return "\n".join(formatted_cv)

def generate_improved_cv_document(resume_text, formatted_content, output_format="docx", experience_years=0):
    """
    Génère un document CV amélioré dans le format spécifié
    
    Args:
        resume_text (str): Le texte original du CV (pour référence)
        formatted_content (str): Le contenu reformaté du CV
        output_format (str): Format de sortie ("docx", "pdf", "txt")
        experience_years (int): Années d'expérience pour le nom du fichier
        
    Returns:
        tuple: (bytes_io, filename) le document généré et son nom
    """
    # Créer un nom de fichier avec les années d'expérience
    exp_suffix = f"_{experience_years}ans" if experience_years > 0 else ""
    filename = f"CV_Amélioré{exp_suffix}_{datetime.now().strftime('%Y%m%d')}"
    
    if output_format == "txt":
        # Fichier texte simple
        bytes_io = io.BytesIO()
        bytes_io.write(formatted_content.encode('utf-8'))
        bytes_io.seek(0)
        return bytes_io, f"{filename}.txt"
    
    elif output_format == "docx":
        # Générer un document DOCX
        try:
            import docx
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = docx.Document()
            
            # Style de document
            styles = doc.styles
            
            # Parcourir le contenu formaté et l'ajouter au document
            for line in formatted_content.split('\n'):
                if line.strip() and all(c.isupper() for c in line.strip() if c.isalpha()):
                    # Titre de section
                    p = doc.add_paragraph()
                    run = p.add_run(line.strip())
                    run.bold = True
                    run.font.size = Pt(14)
                elif line.strip() and line.startswith('-----'):
                    # Ligne de séparation
                    p = doc.add_paragraph()
                elif line.strip() and line.startswith('•'):
                    # Élément de liste
                    p = doc.add_paragraph()
                    p.add_run(line.strip())
                    p.paragraph_format.left_indent = Pt(12)
                elif line.strip():
                    # Vérifier si c'est une ligne contenant l'expérience
                    experience_pattern = r'(\d+)\s*(?:an(?:s|née(?:s)?)|year(?:s)?)\s*(?:d\'?expérience|of\s+experience)'
                    exp_match = re.search(experience_pattern, line)
                    
                    if exp_match:
                        # Mettre en évidence les années d'expérience
                        p = doc.add_paragraph()
                        before_exp = line[:exp_match.start()]
                        experience_text = exp_match.group(0)
                        after_exp = line[exp_match.end():]
                        
                        if before_exp:
                            p.add_run(before_exp)
                        
                        # Mettre l'expérience en gras et en couleur
                        exp_run = p.add_run(experience_text)
                        exp_run.bold = True
                        exp_run.font.color.rgb = RGBColor(0, 128, 0)  # Vert pour l'expérience
                        
                        if after_exp:
                            p.add_run(after_exp)
                    else:
                        # Texte normal
                        p = doc.add_paragraph()
                        p.add_run(line.strip())
                else:
                    # Ligne vide
                    doc.add_paragraph()
            
            # Sauvegarder le document
            bytes_io = io.BytesIO()
            doc.save(bytes_io)
            bytes_io.seek(0)
            return bytes_io, f"{filename}.docx"
            
        except ImportError:
            # Fallback en cas d'absence de python-docx
            bytes_io = io.BytesIO()
            bytes_io.write(formatted_content.encode('utf-8'))
            bytes_io.seek(0)
            return bytes_io, f"{filename}.txt"
    elif output_format == "pdf":
        # Générer un PDF
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.enums import TA_LEFT, TA_CENTER
            
            bytes_io = io.BytesIO()
            doc = SimpleDocTemplate(bytes_io, pagesize=letter)
            styles = getSampleStyleSheet()
            
            heading_style = styles['Heading2']
            normal_style = styles['Normal']
            
            # Liste des éléments à ajouter au document
            elements = []
            
            # Parcourir le contenu formaté et l'ajouter au document
            for line in formatted_content.split('\n'):
                if line.strip() and all(c.isupper() for c in line.strip() if c.isalpha()):
                    # Titre de section
                    elements.append(Paragraph(line.strip(), heading_style))
                elif line.strip() and line.startswith('-----'):
                    # Ligne de séparation
                    elements.append(Spacer(1, 2))
                elif line.strip() and line.startswith('•'):
                    # Élément de liste - utiliser le style Normal avec indentation
                    para = Paragraph(line.strip(), normal_style)
                    para.leftIndent = 20
                    elements.append(para)
                elif line.strip():
                    # Texte normal
                    elements.append(Paragraph(line.strip(), normal_style))
                else:
                    # Ligne vide
                    elements.append(Spacer(1, 12))
            
            # Générer le PDF
            doc.build(elements)
            bytes_io.seek(0)
            return bytes_io, f"{filename}.pdf"
            
        except ImportError:
            # Fallback en cas d'absence de reportlab
            bytes_io = io.BytesIO()
            bytes_io.write(formatted_content.encode('utf-8'))
            bytes_io.seek(0)
            return bytes_io, f"{filename}.txt"
    
    else:
        # Format par défaut - txt
        bytes_io = io.BytesIO()
        bytes_io.write(formatted_content.encode('utf-8'))
        bytes_io.seek(0)
        return bytes_io, f"{filename}.txt"

def improve_cv(file, improvement_target, job_description=None, output_format="docx"):
    """
    Analyse un CV et génère une version améliorée selon la cible spécifiée
    
    Args:
        file: Le fichier CV téléchargé
        improvement_target (str): La cible d'amélioration ("job_offer", "american", "canadian", "linkedin")
        job_description (str, optional): Description du poste si improvement_target est "job_offer"
        output_format (str): Format du document de sortie ("docx", "pdf", "txt")
        
    Returns:
        dict: Résultat contenant l'analyse, les recommandations et le lien vers le document amélioré
    """
    try:
        # Vérifier si le fichier est valide
        if not file or not allowed_file(file.filename):
            return {
                "error": "Format de fichier non pris en charge. Veuillez télécharger un fichier PDF, DOCX ou TXT."
            }
        
        # Extraire le texte du CV
        resume_text = extract_text_from_file(file)
        
        # Prétraiter le texte
        processed_text = preprocess_resume_text(resume_text)
        
        # Extraire les années d'expérience pour le contexte
        experience_years = 0
        experience_patterns = [
            # Français
            r'(\d+)[+\-]?\s*(?:an(?:s|née(?:s)?)?\s+d\'?(?:expérience))',
            r'expérience\s+(?:de|:)?\s*(\d+)[+\-]?\s*an(?:s|née(?:s)?)?',
            # Anglais
            r'(\d+)[+\-]?\s*(?:year(?:s)?\s+(?:of\s+)?experience)',
            r'experience\s+(?:of|:)?\s*(\d+)[+\-]?\s*year(?:s)?'
        ]
        
        for pattern in experience_patterns:
            matches = re.findall(pattern, processed_text.lower())
            if matches:
                try:
                    experience_years = max(experience_years, int(matches[0]))
                    break
                except (ValueError, TypeError, IndexError):
                    continue
        
        # Vérifier la cible d'amélioration
        valid_targets = ["job_offer", "american", "canadian", "linkedin", "general"]
        if improvement_target not in valid_targets:
            improvement_target = "general"
        
        # Obtenir le contexte supplémentaire pour le prompt avec les années d'expérience
        additional_context = get_additional_context_for_improvement(
            improvement_target, 
            job_description,
            experience_years
        )
        
        # Tenter d'analyser le CV avec le LLM
        try:
            prompt = CV_IMPROVEMENT_PROMPT.format(
                improvement_target=improvement_target,
                resume_text=processed_text[:7000],  # Limiter la taille pour éviter de dépasser le contexte
                additional_context=additional_context
            )
            
            llm_response = get_llm_response(prompt)
            
            try:
                # Tenter de parser la réponse JSON
                analysis = json.loads(llm_response)
                
                # Ajouter des métadonnées et les années d'expérience
                analysis['metadata'] = {
                    'filename': secure_filename(file.filename),
                    'analysis_timestamp': datetime.now().isoformat(),
                    'improvement_target': improvement_target,
                    'analysis_method': 'llm'
                }
                
                # Ajouter les années d'expérience si non incluses
                if 'experience_years' not in analysis:
                    analysis['experience_years'] = experience_years
                
                # Utiliser le LLM pour reformater le contenu du CV
                format_prompt = f"""
                En te basant sur l'analyse suivante d'un CV:
                ```
                {json.dumps(analysis, ensure_ascii=False)}
                ```
                
                Et sur le contenu original du CV:
                ```
                {processed_text[:5000]}
                ```
                
                Reformate complètement ce CV pour le standard {improvement_target}.
                {"Optimise-le pour l'offre d'emploi suivante: " + job_description if improvement_target == "job_offer" and job_description else ""}
                
                Le candidat a {experience_years} ans d'expérience professionnelle. Assure-toi que cette information est clairement mise en valeur dans le résumé et si pertinent, dans l'en-tête du CV.
                
                Inclus toutes les sections pertinentes et restructure le contenu pour maximiser l'impact selon les recommandations.
                Ajoute des réalisations quantifiées et des résultats mesurables pour chaque expérience.
                
                Retourne uniquement le contenu formaté du CV, sans explications ni commentaires supplémentaires.
                """
                
                try:
                    formatted_content = get_llm_response(format_prompt)
                    
                    # Générer le document CV amélioré
                    document_io, filename = generate_improved_cv_document(
                        processed_text, 
                        formatted_content, 
                        output_format,
                        analysis.get('experience_years', experience_years)
                    )
                    
                    # Ajouter le contenu formaté à l'analyse
                    analysis['formatted_content'] = formatted_content
                    analysis['output_document'] = {
                        'filename': filename,
                        'format': output_format
                    }
                    
                    # Retourner le résultat avec les années d'expérience mises en évidence
                    return {
                        'analysis': analysis,
                        'document': (document_io, filename),
                        'experience_years': analysis.get('experience_years', experience_years)
                    }
                    
                except Exception as format_error:
                    # En cas d'erreur avec le formatage LLM, utiliser la méthode de secours
                    print(f"Erreur de formatage LLM: {str(format_error)}, utilisation de la méthode de secours")
                    formatted_content = format_cv_content_no_llm(
                        processed_text, 
                        improvement_target, 
                        job_description, 
                        analysis
                    )
                    
                    # Générer le document CV amélioré
                    document_io, filename = generate_improved_cv_document(
                        processed_text, 
                        formatted_content, 
                        output_format,
                        analysis.get('experience_years', experience_years)
                    )
                    
                    # Ajouter le contenu formaté à l'analyse
                    analysis['formatted_content'] = formatted_content
                    analysis['output_document'] = {
                        'filename': filename,
                        'format': output_format
                    }
                    analysis['metadata']['formatting_method'] = 'fallback'
                    
                    # Retourner le résultat
                    return {
                        'analysis': analysis,
                        'document': (document_io, filename),
                        'experience_years': analysis.get('experience_years', experience_years)
                    }
                
            except json.JSONDecodeError:
                # Si la réponse n'est pas au format JSON valide, utiliser la méthode de secours
                print(f"Erreur de décodage JSON dans la réponse LLM, utilisation de la méthode de secours")
                fallback_analysis = improve_cv_no_llm(processed_text, improvement_target, job_description)
                
                # Ajouter des métadonnées
                fallback_analysis['metadata'] = {
                    'filename': secure_filename(file.filename),
                    'analysis_timestamp': datetime.now().isoformat(),
                    'improvement_target': improvement_target,
                    'analysis_method': 'fallback',
                    'error': 'JSON invalide dans la réponse LLM'
                }
                
                # Formater le contenu
                formatted_content = format_cv_content_no_llm(
                    processed_text, 
                    improvement_target, 
                    job_description, 
                    fallback_analysis
                )
                
                # Générer le document CV amélioré
                document_io, filename = generate_improved_cv_document(
                    processed_text, 
                    formatted_content, 
                    output_format,
                    fallback_analysis.get('experience_years', experience_years)
                )
                
                # Ajouter le contenu formaté à l'analyse
                fallback_analysis['formatted_content'] = formatted_content
                fallback_analysis['output_document'] = {
                    'filename': filename,
                    'format': output_format
                }
                
                # Retourner le résultat
                return {
                    'analysis': fallback_analysis,
                    'document': (document_io, filename),
                    'experience_years': fallback_analysis.get('experience_years', experience_years)
                }
                
        except Exception as llm_error:
            # En cas d'erreur avec le LLM, utiliser la méthode de secours
            print(f"Erreur LLM: {str(llm_error)}, utilisation de la méthode de secours")
            fallback_analysis = improve_cv_no_llm(processed_text, improvement_target, job_description)
            
            # Ajouter des métadonnées
            fallback_analysis['metadata'] = {
                'filename': secure_filename(file.filename),
                'analysis_timestamp': datetime.now().isoformat(),
                'improvement_target': improvement_target,
                'analysis_method': 'fallback',
                'error': f'Erreur LLM: {str(llm_error)}'
            }
            
            # Formater le contenu
            formatted_content = format_cv_content_no_llm(
                processed_text, 
                improvement_target, 
                job_description, 
                fallback_analysis
            )
            
            # Générer le document CV amélioré
            document_io, filename = generate_improved_cv_document(
                processed_text, 
                formatted_content, 
                output_format,
                fallback_analysis.get('experience_years', experience_years)
            )
            
            # Ajouter le contenu formaté à l'analyse
            fallback_analysis['formatted_content'] = formatted_content
            fallback_analysis['output_document'] = {
                'filename': filename,
                'format': output_format
            }
            
            # Retourner le résultat
            return {
                'analysis': fallback_analysis,
                'document': (document_io, filename),
                'experience_years': fallback_analysis.get('experience_years', experience_years)
            }
            
    except Exception as e:
        # Capturer toute autre erreur
        return {
            "error": f"Erreur lors de l'amélioration du CV: {str(e)}",
            "analysis_method": "failed"
        }
    
def detect_language(text):
    """
    Détecte la langue principale du texte (français ou anglais)
    
    Args:
        text (str): Texte à analyser
        
    Returns:
        str: 'fr' pour français, 'en' pour anglais, 'unknown' si indéterminé
    """
    # Mots fréquents en français
    fr_words = set(['de', 'la', 'et', 'le', 'du', 'en', 'un', 'une', 'des', 'pour', 'avec', 'par', 'sur', 'dans', 'au', 'à'])
    
    # Mots fréquents en anglais
    en_words = set(['the', 'of', 'and', 'to', 'in', 'a', 'for', 'with', 'on', 'at', 'by', 'from', 'as', 'an', 'is', 'are'])
    
    # Préparer le texte
    words = re.findall(r'\b\w+\b', text.lower())
    
    # Compter les occurrences
    fr_count = sum(1 for word in words if word in fr_words)
    en_count = sum(1 for word in words if word in en_words)
    
    # Sections courantes en français
    fr_sections = sum(1 for line in text.lower().split('\n') if any(s in line for s in ['expérience', 'compétences', 'formation', 'éducation', 'langues']))
    
    # Sections courantes en anglais
    en_sections = sum(1 for line in text.lower().split('\n') if any(s in line for s in ['experience', 'skills', 'education', 'languages', 'certifications']))
    
    # Déterminer la langue
    fr_score = fr_count + (fr_sections * 3)  # Donner un poids plus important aux sections
    en_score = en_count + (en_sections * 3)
    
    if fr_score > en_score * 1.2:  # Si le français est clairement dominant
        return 'fr'
    elif en_score > fr_score * 1.2:  # Si l'anglais est clairement dominant
        return 'en'
    else:
        # Si pas de différence claire, vérifier des phrases plus spécifiques
        if 'years of experience' in text.lower() or 'work experience' in text.lower():
            return 'en'
        elif "années d'expérience" in text.lower() or "expérience professionnelle" in text.lower():
            return 'fr'
        else:
            return 'unknown' 
        
